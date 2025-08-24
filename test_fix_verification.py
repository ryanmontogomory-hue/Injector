from docx import Document
from document_processor import DocumentProcessor

def test_fix_verification():
    print("Testing the fix for point duplication issue...")
    
    # Create a test document with multiple projects
    doc = Document()
    doc.add_heading('Test Resume', 0)
    
    # Project 1
    doc.add_paragraph('Company A | Jan 2020 - Present')
    doc.add_paragraph('E-commerce Platform')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Existing responsibility 1')
    doc.add_paragraph('- Existing responsibility 2')
    
    doc.add_paragraph('')  # Empty line
    
    # Project 2
    doc.add_paragraph('Company B | Feb 2018 - Dec 2019')
    doc.add_paragraph('Data Analytics Dashboard')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Existing responsibility A')
    doc.add_paragraph('- Existing responsibility B')
    
    doc.add_paragraph('')  # Empty line
    
    # Project 3
    doc.add_paragraph('Company C | Sep 2016 - Jan 2018')
    doc.add_paragraph('Customer Management System')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Existing responsibility X')
    doc.add_paragraph('- Existing responsibility Y')
    
    print("Document created with 3 projects")
    
    # Create document processor
    processor = DocumentProcessor()
    
    # Detect projects
    projects_data = processor.project_detector.find_projects_and_responsibilities(doc)
    print(f"Detected projects: {len(projects_data)}")
    
    # Convert to structured format
    projects = []
    for i, (title, start_idx, end_idx) in enumerate(projects_data):
        projects.append({
            'title': title,
            'index': i,
            'responsibilities_start': start_idx,
            'responsibilities_end': end_idx
        })
    
    # Create test tech stacks
    tech_stacks = {
        'React': ['Built responsive UI', 'Used hooks', 'Created components'],
        'MongoDB': ['Created database schema', 'Optimized queries'],
        'Node.js': ['Built REST API', 'Used JWT'],
        'Python': ['Built data processing pipelines', 'Implemented ML models']
    }
    
    print(f"Tech stacks: {len(tech_stacks)} stacks with {sum(len(points) for points in tech_stacks.values())} points")
    
    # Distribute points using the fixed method
    distribution_result = processor.point_distributor.distribute_points_to_projects(projects, tech_stacks)
    
    print(f"\nDistribution result success: {distribution_result['success']}")
    if not distribution_result['success']:
        print(f"Error: {distribution_result['error']}")
        return
    
    print(f"Points added: {distribution_result['points_added']}")
    print(f"Projects used: {distribution_result['projects_used']}")
    
    # Print distribution details
    print("\nDistribution details:")
    for project_title, project_info in distribution_result['distribution'].items():
        print(f"  {project_title}: {project_info['total_points']} points")
    
    # Process the document using the fixed method
    print("\nProcessing document with fixed method...")
    
    # Simulate the fixed process_document method
    total_added = 0
    # Sort projects by insertion point to process them in order
    sorted_projects = sorted(distribution_result['distribution'].items(), 
                           key=lambda x: x[1]['insertion_point'])
    
    # Keep track of how many paragraphs we've added to adjust insertion points
    paragraph_offset = 0
    
    for project_title, project_info in sorted_projects:
        # Adjust insertion point based on previous additions
        adjusted_project_info = project_info.copy()
        adjusted_project_info['insertion_point'] += paragraph_offset
        if 'responsibilities_end' in adjusted_project_info:
            adjusted_project_info['responsibilities_end'] += paragraph_offset
        
        added = processor.add_points_to_project(doc, adjusted_project_info)
        print(f"Added {added} points to {project_title}")
        total_added += added
        
        # Update the offset for subsequent projects
        paragraph_offset += added
    
    print(f"\nTotal points added: {total_added}")
    
    # Verify that each project has only its designated points
    print("\nVerifying point distribution...")
    
    # Count points in each project section
    project_sections = []
    current_project = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if '|' in text and 'Company' in text:
            current_project = text.split('|')[0].strip()
            project_sections.append({'name': current_project, 'start': i, 'points': []})
        elif text.startswith('-') and current_project and len(project_sections) > 0:
            # Check if this is a newly added point (not an existing responsibility)
            existing_responsibilities = [
                'Existing responsibility 1', 'Existing responsibility 2',
                'Existing responsibility A', 'Existing responsibility B',
                'Existing responsibility X', 'Existing responsibility Y'
            ]
            if not any(existing in text for existing in existing_responsibilities):
                project_sections[-1]['points'].append(text)
    
    # Print verification results
    for section in project_sections:
        print(f"Project {section['name']}: {len(section['points'])} added points")
        for point in section['points']:
            print(f"  {point}")
    
    # Check for duplicates
    all_added_points = []
    for section in project_sections:
        all_added_points.extend(section['points'])
    
    # Check for duplicates
    seen = set()
    duplicates = []
    for point in all_added_points:
        if point in seen:
            duplicates.append(point)
        seen.add(point)
    
    if duplicates:
        print(f"\nERROR: Found duplicate points: {duplicates}")
        return False
    else:
        print("\nSUCCESS: No duplicate points found!")
        return True

# Run the test
if __name__ == "__main__":
    success = test_fix_verification()
    if success:
        print("\n[PASS] Fix verification PASSED!")
    else:
        print("\n[FAIL] Fix verification FAILED!")