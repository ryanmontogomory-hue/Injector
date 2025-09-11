"""
Enhanced error recovery system for Resume Customizer application.
Provides robust error handling, automatic retry mechanisms, and graceful degradation.
"""

import gc
import time
import traceback
import threading
from typing import Dict, Any, Optional, Callable, List, Union, Type
from dataclasses import dataclass, field
from functools import wraps
from enum import Enum
import psutil
from io import BytesIO

from utilities.logger import get_logger
from utilities.structured_logger import get_structured_logger
from .error_handling_enhanced import ErrorHandler, ErrorContext, ErrorSeverity
from monitoring.circuit_breaker import CircuitBreakerError, get_circuit_breaker, CircuitBreakerConfig

logger = get_logger()
structured_logger = get_structured_logger("error_recovery")


class RecoveryStrategy(Enum):
    """Error recovery strategies."""
    RETRY = "retry"
    FALLBACK = "fallback"
    DEGRADE = "degrade"
    FAIL_FAST = "fail_fast"
    CIRCUIT_BREAK = "circuit_break"


@dataclass
class RetryConfig:
    """Configuration for retry mechanism."""
    max_attempts: int = 3
    backoff_factor: float = 2.0
    initial_delay: float = 1.0
    max_delay: float = 60.0
    exponential_backoff: bool = True
    jitter: bool = True
    retryable_exceptions: tuple = field(default_factory=lambda: (Exception,))


@dataclass
class RecoveryResult:
    """Result of error recovery attempt."""
    success: bool
    result: Any = None
    error: Optional[Exception] = None
    strategy_used: Optional[RecoveryStrategy] = None
    attempts_made: int = 0
    recovery_time: float = 0.0
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentCorruptionError(Exception):
    """Raised when document appears to be corrupted."""
    pass


class MemoryPressureError(Exception):
    """Raised when system is under memory pressure."""
    pass


def retry_with_backoff(config: Optional[RetryConfig] = None):
    """
    Decorator for retry with exponential backoff and jitter.
    
    Args:
        config: Retry configuration
    """
    if config is None:
        config = RetryConfig()
    
    def decorator(func: Callable) -> Callable:
        @wraps(func)
        def wrapper(*args, **kwargs):
            last_exception = None
            delay = config.initial_delay
            
            for attempt in range(config.max_attempts):
                try:
                    start_time = time.time()
                    result = func(*args, **kwargs)
                    
                    if attempt > 0:  # Log successful retry
                        structured_logger.info(
                            f"Function {func.__name__} succeeded on attempt {attempt + 1}",
                            operation="retry_success",
                            attempts=attempt + 1,
                            duration=time.time() - start_time
                        )
                    
                    return result
                    
                except config.retryable_exceptions as e:
                    last_exception = e
                    
                    if attempt < config.max_attempts - 1:  # Not the last attempt
                        # Calculate delay with jitter
                        actual_delay = delay
                        if config.jitter:
                            import random
                            actual_delay *= (0.5 + random.random() * 0.5)  # 50-100% of delay
                        
                        structured_logger.warning(
                            f"Function {func.__name__} failed on attempt {attempt + 1}, "
                            f"retrying in {actual_delay:.2f}s",
                            operation="retry_attempt",
                            attempt=attempt + 1,
                            max_attempts=config.max_attempts,
                            delay=actual_delay,
                            error=str(e)
                        )
                        
                        time.sleep(actual_delay)
                        
                        # Calculate next delay
                        if config.exponential_backoff:
                            delay = min(delay * config.backoff_factor, config.max_delay)
                    else:
                        structured_logger.error(
                            f"Function {func.__name__} failed after {config.max_attempts} attempts",
                            operation="retry_exhausted",
                            attempts=config.max_attempts,
                            final_error=str(e)
                        )
                
                except Exception as e:
                    # Non-retryable exception
                    structured_logger.error(
                        f"Function {func.__name__} failed with non-retryable exception",
                        operation="retry_non_retryable",
                        error=str(e),
                        exception_type=type(e).__name__
                    )
                    raise
            
            # All retries exhausted
            raise last_exception
        
        return wrapper
    return decorator


class EnhancedErrorRecovery:
    """Enhanced error recovery with multiple strategies."""
    
    def __init__(self):
        self.error_handler = ErrorHandler()
        self.recovery_stats = {
            'total_errors': 0,
            'successful_recoveries': 0,
            'failed_recoveries': 0,
            'strategy_usage': {strategy.value: 0 for strategy in RecoveryStrategy}
        }
        self._stats_lock = threading.Lock()
    
    def recover_with_strategy(
        self,
        func: Callable,
        strategy: RecoveryStrategy,
        fallback_func: Optional[Callable] = None,
        context: Optional[ErrorContext] = None,
        *args,
        **kwargs
    ) -> RecoveryResult:
        """
        Execute function with specified recovery strategy.
        
        Args:
            func: Primary function to execute
            strategy: Recovery strategy to use
            fallback_func: Fallback function for FALLBACK strategy
            context: Error context for logging
            *args, **kwargs: Arguments for the function
        """
        start_time = time.time()
        attempts = 0
        
        with self._stats_lock:
            self.recovery_stats['total_errors'] += 1
            self.recovery_stats['strategy_usage'][strategy.value] += 1
        
        try:
            if strategy == RecoveryStrategy.RETRY:
                return self._retry_strategy(func, context, *args, **kwargs)
            
            elif strategy == RecoveryStrategy.FALLBACK:
                return self._fallback_strategy(func, fallback_func, context, *args, **kwargs)
            
            elif strategy == RecoveryStrategy.DEGRADE:
                return self._degrade_strategy(func, context, *args, **kwargs)
            
            elif strategy == RecoveryStrategy.FAIL_FAST:
                return self._fail_fast_strategy(func, context, *args, **kwargs)
            
            elif strategy == RecoveryStrategy.CIRCUIT_BREAK:
                return self._circuit_break_strategy(func, context, *args, **kwargs)
            
            else:
                raise ValueError(f"Unknown recovery strategy: {strategy}")
        
        except Exception as e:
            with self._stats_lock:
                self.recovery_stats['failed_recoveries'] += 1
            
            return RecoveryResult(
                success=False,
                error=e,
                strategy_used=strategy,
                attempts_made=attempts,
                recovery_time=time.time() - start_time,
                metadata={'strategy': strategy.value}
            )
    
    def _retry_strategy(self, func: Callable, context: Optional[ErrorContext], *args, **kwargs) -> RecoveryResult:
        """Retry strategy with exponential backoff."""
        config = RetryConfig(
            max_attempts=3,
            backoff_factor=2.0,
            retryable_exceptions=(
                ConnectionError, TimeoutError, MemoryError, 
                IOError, DocumentCorruptionError
            )
        )
        
        decorated_func = retry_with_backoff(config)(func)
        
        try:
            result = decorated_func(*args, **kwargs)
            with self._stats_lock:
                self.recovery_stats['successful_recoveries'] += 1
            
            return RecoveryResult(
                success=True,
                result=result,
                strategy_used=RecoveryStrategy.RETRY,
                attempts_made=config.max_attempts
            )
        
        except Exception as e:
            return RecoveryResult(
                success=False,
                error=e,
                strategy_used=RecoveryStrategy.RETRY
            )
    
    def _fallback_strategy(
        self, 
        func: Callable, 
        fallback_func: Optional[Callable],
        context: Optional[ErrorContext],
        *args, 
        **kwargs
    ) -> RecoveryResult:
        """Fallback strategy - try primary, then fallback function."""
        try:
            result = func(*args, **kwargs)
            return RecoveryResult(
                success=True,
                result=result,
                strategy_used=RecoveryStrategy.FALLBACK,
                metadata={'used_primary': True}
            )
        
        except Exception as primary_error:
            if fallback_func is None:
                return RecoveryResult(
                    success=False,
                    error=primary_error,
                    strategy_used=RecoveryStrategy.FALLBACK,
                    metadata={'fallback_available': False}
                )
            
            try:
                structured_logger.info(
                    "Primary function failed, attempting fallback",
                    operation="fallback_attempt",
                    primary_error=str(primary_error)
                )
                
                result = fallback_func(*args, **kwargs)
                
                with self._stats_lock:
                    self.recovery_stats['successful_recoveries'] += 1
                
                return RecoveryResult(
                    success=True,
                    result=result,
                    strategy_used=RecoveryStrategy.FALLBACK,
                    metadata={'used_fallback': True, 'primary_error': str(primary_error)}
                )
            
            except Exception as fallback_error:
                return RecoveryResult(
                    success=False,
                    error=fallback_error,
                    strategy_used=RecoveryStrategy.FALLBACK,
                    metadata={
                        'primary_error': str(primary_error),
                        'fallback_error': str(fallback_error)
                    }
                )
    
    def _degrade_strategy(self, func: Callable, context: Optional[ErrorContext], *args, **kwargs) -> RecoveryResult:
        """Degrade strategy - return partial or simplified result."""
        try:
            result = func(*args, **kwargs)
            return RecoveryResult(
                success=True,
                result=result,
                strategy_used=RecoveryStrategy.DEGRADE
            )
        
        except Exception as e:
            # Attempt to provide a degraded response
            degraded_result = self._create_degraded_result(func.__name__, e, *args, **kwargs)
            
            if degraded_result is not None:
                with self._stats_lock:
                    self.recovery_stats['successful_recoveries'] += 1
                
                return RecoveryResult(
                    success=True,
                    result=degraded_result,
                    strategy_used=RecoveryStrategy.DEGRADE,
                    metadata={'degraded': True, 'original_error': str(e)}
                )
            
            return RecoveryResult(
                success=False,
                error=e,
                strategy_used=RecoveryStrategy.DEGRADE
            )
    
    def _fail_fast_strategy(self, func: Callable, context: Optional[ErrorContext], *args, **kwargs) -> RecoveryResult:
        """Fail fast strategy - no recovery, immediate failure."""
        try:
            result = func(*args, **kwargs)
            return RecoveryResult(
                success=True,
                result=result,
                strategy_used=RecoveryStrategy.FAIL_FAST
            )
        
        except Exception as e:
            return RecoveryResult(
                success=False,
                error=e,
                strategy_used=RecoveryStrategy.FAIL_FAST,
                metadata={'fail_fast': True}
            )
    
    def _circuit_break_strategy(self, func: Callable, context: Optional[ErrorContext], *args, **kwargs) -> RecoveryResult:
        """Circuit breaker strategy."""
        circuit_name = func.__name__
        circuit_breaker = get_circuit_breaker(
            circuit_name,
            CircuitBreakerConfig(failure_threshold=5, recovery_timeout=30)
        )
        
        try:
            result = circuit_breaker.call(func, *args, **kwargs)
            return RecoveryResult(
                success=True,
                result=result,
                strategy_used=RecoveryStrategy.CIRCUIT_BREAK
            )
        
        except CircuitBreakerError as e:
            return RecoveryResult(
                success=False,
                error=e,
                strategy_used=RecoveryStrategy.CIRCUIT_BREAK,
                metadata={'circuit_open': True}
            )
        
        except Exception as e:
            return RecoveryResult(
                success=False,
                error=e,
                strategy_used=RecoveryStrategy.CIRCUIT_BREAK
            )
    
    def _create_degraded_result(self, func_name: str, error: Exception, *args, **kwargs) -> Optional[Any]:
        """Create a degraded result when primary function fails."""
        if "process" in func_name.lower():
            # For processing functions, return a minimal result
            return {
                'success': False,
                'error': f'Processing failed: {str(error)}',
                'degraded': True,
                'filename': kwargs.get('filename', 'unknown'),
                'fallback_message': 'The file could not be processed due to an error.'
            }
        
        elif "parse" in func_name.lower():
            # For parsing functions, return empty results
            return [], []
        
        elif "email" in func_name.lower():
            # For email functions, return failure result
            return {
                'success': False,
                'error': f'Email sending failed: {str(error)}',
                'degraded': True
            }
        
        return None
    
    def get_stats(self) -> Dict[str, Any]:
        """Get recovery statistics."""
        with self._stats_lock:
            stats = self.recovery_stats.copy()
        
        if stats['total_errors'] > 0:
            stats['recovery_rate'] = stats['successful_recoveries'] / stats['total_errors']
        else:
            stats['recovery_rate'] = 0
        
        return stats


class RobustResumeProcessor:
    """Resume processor with enhanced error recovery and memory management."""
    
    def __init__(self):
        # Avoid circular import by lazy loading
        self.base_processor = None
        self.recovery_manager = EnhancedErrorRecovery()
        self.memory_threshold_mb = 500
        self.emergency_cleanup_threshold_mb = 800
    
    def _get_base_processor(self):
        """Lazy load the base processor to avoid circular imports."""
        if self.base_processor is None:
            from core.resume_processor import ResumeProcessor
            self.base_processor = ResumeProcessor()
        return self.base_processor
    
    @retry_with_backoff(RetryConfig(
        max_attempts=3,
        retryable_exceptions=(IOError, MemoryError, DocumentCorruptionError)
    ))
    def process_single_resume(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process single resume with enhanced error recovery."""
        # Check memory before processing
        self._check_memory_pressure()
        
        try:
            # Attempt normal processing
            base_processor = self._get_base_processor()
            result = base_processor.process_single_resume(file_data)
            return result
            
        except MemoryError:
            # Try memory-optimized processing
            logger.warning("MemoryError detected, attempting memory-optimized processing")
            return self._process_with_reduced_memory(file_data)
        
        except Exception as e:
            # Check if document might be corrupted
            if self._is_document_corruption_error(e):
                return self._attempt_document_repair(file_data)
            
            # For other errors, use recovery strategy
            base_processor = self._get_base_processor()
            recovery_result = self.recovery_manager.recover_with_strategy(
                base_processor.process_single_resume,
                RecoveryStrategy.DEGRADE,
                self._fallback_processing,
                ErrorContext(
                    operation="process_single_resume",
                    file_name=file_data.get('filename', 'unknown')
                ),
                file_data
            )
            
            if recovery_result.success:
                return recovery_result.result
            else:
                raise recovery_result.error
    
    def _check_memory_pressure(self):
        """Check system memory pressure and take action if needed."""
        try:
            memory = psutil.virtual_memory()
            memory_usage_mb = memory.used / (1024 * 1024)
            
            if memory_usage_mb > self.emergency_cleanup_threshold_mb:
                logger.warning(f"Emergency memory cleanup triggered: {memory_usage_mb:.1f}MB used")
                self._emergency_memory_cleanup()
                raise MemoryPressureError("System under extreme memory pressure")
            
            elif memory_usage_mb > self.memory_threshold_mb:
                logger.info(f"Memory cleanup triggered: {memory_usage_mb:.1f}MB used")
                self._gentle_memory_cleanup()
                
        except Exception as e:
            logger.warning(f"Memory check failed: {e}")
    
    def _gentle_memory_cleanup(self):
        """Perform gentle memory cleanup."""
        try:
            # Clear caches
            from monitoring.performance_cache import get_cache_manager
            cache_manager = get_cache_manager()
            for cache_name in ['processing', 'parsing', 'document']:
                cache = cache_manager.get_cache(cache_name)
                # Clear 50% of cache entries
                if hasattr(cache, '_cache'):
                    items_to_remove = len(cache._cache) // 2
                    for _ in range(items_to_remove):
                        if cache._cache:
                            cache._cache.popitem(last=False)  # Remove oldest
            
            gc.collect()
            logger.info("Gentle memory cleanup completed")
            
        except Exception as e:
            logger.warning(f"Gentle memory cleanup failed: {e}")
    
    def _emergency_memory_cleanup(self):
        """Perform aggressive memory cleanup."""
        try:
            # Clear all caches
            from monitoring.performance_cache import get_cache_manager
            cache_manager = get_cache_manager()
            for cache_name in ['processing', 'parsing', 'document']:
                cache_manager.get_cache(cache_name).clear()
            
            # Force garbage collection multiple times
            for _ in range(3):
                gc.collect()
            
            logger.info("Emergency memory cleanup completed")
            
        except Exception as e:
            logger.error(f"Emergency memory cleanup failed: {e}")
    
    def _process_with_reduced_memory(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process resume with reduced memory usage."""
        try:
            # Clear caches before processing
            self._gentle_memory_cleanup()
            
            # Create a simplified version of processing
            filename = file_data.get('filename', 'unknown')
            
            # Attempt minimal processing
            return {
                'success': True,
                'filename': filename,
                'tech_stacks': ['General'],  # Fallback tech stack
                'points_added': 1,  # Minimal points
                'buffer': self._create_minimal_document(file_data),
                'memory_optimized': True
            }
            
        except Exception as e:
            logger.error(f"Reduced memory processing failed: {e}")
            raise
    
    def _is_document_corruption_error(self, error: Exception) -> bool:
        """Check if error indicates document corruption."""
        error_msg = str(error).lower()
        corruption_indicators = [
            'corrupt', 'invalid', 'malformed', 'bad', 'broken',
            'xml', 'zip', 'format', 'structure'
        ]
        return any(indicator in error_msg for indicator in corruption_indicators)
    
    def _attempt_document_repair(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """Attempt to repair corrupted document."""
        logger.info(f"Attempting document repair for {file_data.get('filename')}")
        
        try:
            # Try to read as raw text if possible
            file_obj = file_data.get('file')
            if file_obj:
                file_obj.seek(0)
                raw_content = file_obj.read()
                
                # Try to extract any readable text
                readable_text = self._extract_readable_text(raw_content)
                if readable_text:
                    return {
                        'success': True,
                        'filename': file_data.get('filename'),
                        'tech_stacks': ['General'],
                        'points_added': 0,
                        'buffer': self._create_text_document(readable_text),
                        'repaired': True,
                        'warning': 'Document was corrupted and has been partially recovered'
                    }
            
            raise DocumentCorruptionError("Unable to repair document")
            
        except Exception as e:
            logger.error(f"Document repair failed: {e}")
            raise DocumentCorruptionError(f"Document repair failed: {e}")
    
    def _extract_readable_text(self, raw_content: bytes) -> Optional[str]:
        """Extract readable text from raw content."""
        try:
            # Try UTF-8 decoding
            text = raw_content.decode('utf-8', errors='ignore')
            # Remove non-printable characters
            readable = ''.join(char for char in text if char.isprintable() or char.isspace())
            return readable if len(readable) > 50 else None
        except Exception:
            return None
    
    def _create_minimal_document(self, file_data: Dict[str, Any]) -> bytes:
        """Create minimal document for fallback processing."""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Resume', 0)
            doc.add_paragraph('This document was processed in reduced memory mode.')
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to create minimal document: {e}")
            return b''
    
    def _create_text_document(self, text: str) -> bytes:
        """Create document from recovered text."""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('Recovered Document', 0)
            doc.add_paragraph(text[:1000])  # Limit text length
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to create text document: {e}")
            return b''
    
    def _fallback_processing(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback processing when main processing fails."""
        filename = file_data.get('filename', 'unknown')
        return {
            'success': False,
            'filename': filename,
            'error': 'Processing failed, fallback mode activated',
            'fallback': True,
            'suggestion': 'Please try uploading the file again or use a different format'
        }


# Global recovery manager instance
_recovery_manager = None


def get_error_recovery_manager() -> EnhancedErrorRecovery:
    """Get global error recovery manager."""
    global _recovery_manager
    if _recovery_manager is None:
        _recovery_manager = EnhancedErrorRecovery()
    return _recovery_manager



