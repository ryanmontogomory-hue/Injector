"""
Simulate the exact Streamlit UI workflow to test bullet consistency
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from bullet_consistency_patch import apply_emergency_patch
from core.text_parser import parse_input_text
from core.document_processor import DocumentProcessor
from docx import Document

# Apply the patch
apply_emergency_patch()

def simulate_ui_workflow():
    """Simulate the exact workflow that happens in the Streamlit UI"""
    
    test_file = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    print("=== SIMULATING STREAMLIT UI WORKFLOW ===")
    print(f"File: {test_file}")
    
    if not os.path.exists(test_file):
        print(f"❌ File not found: {test_file}")
        return False
    
    # Simulate user input - this is what users typically enter in the UI
    user_input = """Java: • Developed scalable enterprise applications using Spring Boot and microservices architecture • Implemented secure REST APIs with JWT authentication and OAuth2 integration • Built high-performance backend systems with Java 8+ features and multithreading
Python: • Created data processing pipelines using pandas and NumPy for large dataset analysis • Developed machine learning models with scikit-learn and TensorFlow • Built automated testing frameworks using pytest and continuous integration
AWS: • Deployed containerized applications using Docker and Kubernetes on EKS • Implemented serverless architectures with Lambda functions and API Gateway • Managed cloud infrastructure with CloudFormation and Terraform
Database: • Designed and optimized relational database schemas in PostgreSQL and MySQL • Implemented NoSQL solutions using MongoDB for flexible document storage • Built data warehousing solutions with ETL processes and performance tuning"""
    
    print(f"\n1. USER INPUT SIMULATION:")
    print("User enters tech stacks with bullet symbols (•):")
    print(user_input[:200] + "...")
    
    try:
        print(f"\n2. TEXT PARSING SIMULATION:")
        # This simulates what happens when the UI parses user input
        selected_points, tech_stacks_used = parse_input_text(user_input)
        
        print(f"   Parser returned {len(selected_points)} points from {len(tech_stacks_used)} tech stacks")
        print(f"   Tech stacks found: {tech_stacks_used}")
        
        # Convert to the format expected by document processor
        # For this test, we'll simulate distributing points across detected tech stacks
        tech_stacks = {}
        if tech_stacks_used and selected_points:
            # Simple distribution: assign points to the first tech stack found
            tech_name = tech_stacks_used[0] if tech_stacks_used else 'General'
            tech_stacks[tech_name] = selected_points
        
        print(f"   Final tech stacks for processing: {list(tech_stacks.keys())}")
        for tech, points in tech_stacks.items():
            print(f"   {tech}: {len(points)} points")
            for point in points[:2]:  # Show first 2 points
                print(f"     Sample: {point[:60]}...")
        
        print(f"\n3. DOCUMENT PROCESSING SIMULATION:")
        
        # Load the file content (simulating file upload)
        with open(test_file, 'rb') as f:
            file_content = f.read()
        
        # Create the exact data structure used by the UI
        file_data = {
            'file_content': file_content,
            'filename': 'Resume Format 1.docx',
            'tech_stacks': tech_stacks
        }
        
        # Process with DocumentProcessor (exactly as UI does)
        processor = DocumentProcessor()
        result = processor.process_document(file_data)
        
        if result.get('success'):
            print("   ✅ Document processing succeeded!")
            print(f"   - Points added: {result.get('points_added', 0)}")
            
            # Save the result (simulating download)
            output_file = test_file.replace('.docx', '_ui_simulation_output.docx')
            with open(output_file, 'wb') as f:
                f.write(result['modified_content'])
            
            print(f"   - Output saved: {output_file}")
            
            # Analyze the final result
            print(f"\n4. FINAL CONSISTENCY CHECK:")
            
            output_doc = Document(output_file)
            formatter = processor.formatter
            
            # Count all bullets and their markers
            bullet_markers = []
            dash_count = 0
            bullet_count = 0
            
            for i, paragraph in enumerate(output_doc.paragraphs):
                text = paragraph.text.strip()
                if formatter._is_bullet_point(text):
                    marker = formatter._extract_bullet_marker(text)
                    bullet_markers.append(marker)
                    if marker == '-':
                        dash_count += 1
                    elif marker == '•':
                        bullet_count += 1
            
            print(f"   Total bullet points: {len(bullet_markers)}")
            print(f"   Dash markers (-): {dash_count}")
            print(f"   Bullet symbols (•): {bullet_count}")
            
            unique_markers = set(bullet_markers)
            print(f"   Unique markers: {unique_markers}")
            
            if len(unique_markers) == 1:
                marker = list(unique_markers)[0]
                print(f"   ✅ PERFECT CONSISTENCY! All bullets use: '{marker}'")
                return True
            else:
                print(f"   ❌ INCONSISTENCY DETECTED!")
                print(f"   Mixed markers found: {unique_markers}")
                
                # Show examples of inconsistency
                print(f"\n   INCONSISTENCY EXAMPLES:")
                for i, paragraph in enumerate(output_doc.paragraphs[:20]):  # Check first 20
                    text = paragraph.text.strip()
                    if formatter._is_bullet_point(text):
                        marker = formatter._extract_bullet_marker(text)
                        print(f"     [{i:2d}] '{marker}' -> {text[:50]}...")
                
                return False
        else:
            print(f"   ❌ Document processing failed: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"❌ UI simulation failed: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = simulate_ui_workflow()
    
    print(f"\n=== UI SIMULATION RESULT ===")
    if success:
        print("🎉 UI SIMULATION PASSED! Bullet consistency works in UI workflow.")
        print("If you're still seeing issues, try:")
        print("1. Hard refresh your browser (Ctrl+F5)")
        print("2. Clear browser cache")
        print("3. Restart the Streamlit server completely")
    else:
        print("❌ UI SIMULATION FAILED! Issue confirmed in UI workflow.")
        print("The patch needs further refinement.")
    
    exit(0 if success else 1)