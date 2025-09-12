"""
Integration test for bullet formatting consistency
Tests the full application workflow to ensure bullet markers are consistent
"""

import sys
import os
import shutil
from io import BytesIO
from docx import Document

# Add the current directory to Python path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from core.document_processor import DocumentProcessor, get_document_processor
from formatters.bullet_formatter import BulletFormatter

def test_full_integration():
    """Test the complete bullet formatting integration"""
    
    print("=== BULLET FORMATTING INTEGRATION TEST ===")
    
    # Test file path
    test_file = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    if not os.path.exists(test_file):
        print(f"❌ Test file not found: {test_file}")
        return False
    
    # Backup the original file
    backup_file = test_file.replace('.docx', '_backup.docx')
    shutil.copy2(test_file, backup_file)
    print(f"✅ Backup created: {backup_file}")
    
    try:
        # Step 1: Test BulletFormatter alone
        print("\n--- Step 1: Testing BulletFormatter bullet detection ---")
        
        doc = Document(test_file)
        formatter = BulletFormatter()
        
        detected_marker = formatter.detect_document_bullet_marker(doc)
        print(f"✅ Detected bullet marker: '{detected_marker}'")
        
        # Step 2: Test DocumentProcessor
        print("\n--- Step 2: Testing DocumentProcessor ---")
        
        processor = get_document_processor()
        
        # Create sample file data
        with open(test_file, 'rb') as f:
            file_content = f.read()
        
        file_data = {
            'file_content': file_content,
            'filename': 'test_resume.docx',
            'tech_stacks': {
                'Python': [
                    'Developed scalable backend systems using Python and Django',
                    'Implemented data processing pipelines with Python pandas',
                    'Created automated testing frameworks using pytest'
                ],
                'Cloud': [
                    'Deployed applications on AWS EC2 and Lambda',
                    'Managed cloud infrastructure using Docker containers',
                    'Implemented CI/CD pipelines with cloud services'
                ],
                'Database': [
                    'Designed relational database schemas using PostgreSQL',
                    'Optimized database queries for improved performance',
                    'Implemented data backup and recovery strategies'
                ]
            }
        }
        
        # Process the document
        result = processor.process_document(file_data)
        
        if not result.get('success'):
            print(f"❌ Document processing failed: {result.get('error')}")
            return False
        
        print(f"✅ Document processing succeeded")
        print(f"   - Points added: {result.get('points_added', 0)}")
        print(f"   - Projects modified: {result.get('projects_modified', 0)}")
        
        # Step 3: Verify bullet consistency
        print("\n--- Step 3: Verifying bullet consistency ---")
        
        # Save the processed document
        output_file = test_file.replace('.docx', '_processed.docx')
        with open(output_file, 'wb') as f:
            f.write(result['modified_content'])
        
        # Load and analyze the processed document
        processed_doc = Document(output_file)
        final_marker = formatter.detect_document_bullet_marker(processed_doc)
        
        print(f"✅ Final document bullet marker: '{final_marker}'")
        
        # Check if markers are consistent
        if detected_marker.strip() == final_marker.strip():
            print("✅ BULLET CONSISTENCY TEST PASSED!")
            print(f"   Original marker: '{detected_marker}'")
            print(f"   Final marker: '{final_marker}'")
            consistency_passed = True
        else:
            print("❌ BULLET CONSISTENCY TEST FAILED!")
            print(f"   Original marker: '{detected_marker}'")
            print(f"   Final marker: '{final_marker}'")
            consistency_passed = False
        
        # Step 4: Count bullet points to verify additions
        print("\n--- Step 4: Verifying bullet point additions ---")
        
        original_bullets = count_bullet_points(Document(test_file), formatter)
        final_bullets = count_bullet_points(processed_doc, formatter)
        
        print(f"✅ Original bullet points: {original_bullets}")
        print(f"✅ Final bullet points: {final_bullets}")
        print(f"✅ Bullets added: {final_bullets - original_bullets}")
        
        additions_passed = final_bullets > original_bullets
        
        if additions_passed:
            print("✅ BULLET ADDITION TEST PASSED!")
        else:
            print("❌ BULLET ADDITION TEST FAILED!")
        
        # Overall result
        overall_passed = consistency_passed and additions_passed
        
        print(f"\n=== INTEGRATION TEST SUMMARY ===")
        print(f"Bullet consistency: {'✅ PASSED' if consistency_passed else '❌ FAILED'}")
        print(f"Bullet additions: {'✅ PASSED' if additions_passed else '❌ FAILED'}")
        print(f"Overall result: {'✅ PASSED' if overall_passed else '❌ FAILED'}")
        
        return overall_passed
        
    except Exception as e:
        print(f"❌ Integration test failed with exception: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        # Restore the original file
        if os.path.exists(backup_file):
            shutil.copy2(backup_file, test_file)
            os.remove(backup_file)
            print(f"✅ Original file restored")

def count_bullet_points(doc: Document, formatter: BulletFormatter) -> int:
    """Count bullet points in a document"""
    count = 0
    for paragraph in doc.paragraphs:
        if formatter._is_bullet_point(paragraph.text):
            count += 1
    return count

def test_bullet_formatter_directly():
    """Test the BulletFormatter directly with the resume file"""
    
    print("\n=== DIRECT BULLET FORMATTER TEST ===")
    
    test_file = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    if not os.path.exists(test_file):
        print(f"❌ Test file not found: {test_file}")
        return False
    
    try:
        doc = Document(test_file)
        formatter = BulletFormatter()
        
        # Test detection
        detected_marker = formatter.detect_document_bullet_marker(doc)
        print(f"✅ Detected marker: '{detected_marker}'")
        
        # Test extraction from specific paragraphs
        bullet_examples = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if formatter._is_bullet_point(text):
                extracted = formatter._extract_bullet_marker(text)
                bullet_examples.append((i, text[:50] + '...', extracted))
                if len(bullet_examples) >= 5:  # Show first 5 examples
                    break
        
        print(f"✅ Found {len(bullet_examples)} bullet examples:")
        for i, (para_idx, text, marker) in enumerate(bullet_examples):
            print(f"   {i+1}. Para {para_idx}: '{marker}' -> {text}")
        
        return True
        
    except Exception as e:
        print(f"❌ Direct test failed: {e}")
        return False

if __name__ == "__main__":
    print("Starting bullet formatting integration tests...\n")
    
    # Run direct formatter test first
    direct_test_passed = test_bullet_formatter_directly()
    
    # Run full integration test
    integration_test_passed = test_full_integration()
    
    print(f"\n=== FINAL RESULTS ===")
    print(f"Direct formatter test: {'✅ PASSED' if direct_test_passed else '❌ FAILED'}")
    print(f"Integration test: {'✅ PASSED' if integration_test_passed else '❌ FAILED'}")
    
    if direct_test_passed and integration_test_passed:
        print("🎉 ALL TESTS PASSED! Bullet formatting is working correctly.")
        exit(0)
    else:
        print("❌ Some tests failed. Please check the output above.")
        exit(1)