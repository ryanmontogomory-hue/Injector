"""
Test the emergency patch with bullet symbols specifically
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from bullet_consistency_patch import apply_emergency_patch, patch_bullet_formatter
from formatters.bullet_formatter import BulletFormatter
from docx import Document
from docx.shared import Inches

# Apply the patch
apply_emergency_patch()

def create_bullet_symbol_test_document():
    """Create a test document with bullet symbols"""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph("Test Resume")
    title.style = 'Title'
    
    # Add Work Experience section (similar to your resume structure)
    work_header = doc.add_paragraph("Work Experience:")
    work_header.style = 'Heading 1'
    
    # Add a project
    project_title = doc.add_paragraph("Test Company | Jan 2020 – Present\nSenior Software Engineer")
    
    # Add Responsibilities section
    resp_header = doc.add_paragraph("Responsibilities:")
    resp_header.style = 'Heading 2'
    
    # Add some bullet points with bullet symbols
    bullet_points = [
        "• Developed web applications using React and Node.js",
        "• Implemented REST APIs with Express.js and MongoDB", 
        "• Deployed applications to AWS using Docker containers",
        "• Optimized database queries for improved performance"
    ]
    
    for point in bullet_points:
        para = doc.add_paragraph(point)
        para.paragraph_format.left_indent = Inches(0.5)
    
    # Save the document
    test_file = r"C:\Users\HP\Downloads\test_bullet_symbols.docx"
    doc.save(test_file)
    print(f"✅ Created test document with bullet symbols: {test_file}")
    
    return test_file

def test_bullet_symbol_detection():
    """Test bullet symbol detection with the patch"""
    
    print("=== TESTING BULLET SYMBOL DETECTION WITH PATCH ===")
    
    # Create test document
    test_file = create_bullet_symbol_test_document()
    
    try:
        # Load the document
        doc = Document(test_file)
        
        # Test the patched formatter
        formatter = BulletFormatter()
        
        print("\n1. Testing bullet symbol detection...")
        detected_marker = formatter.get_document_bullet_marker(doc)
        print(f"✅ Detected marker: '{detected_marker}'")
        
        print("\n2. Testing individual bullet extraction...")
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text.startswith('•'):
                extracted = formatter._extract_bullet_marker(text)
                print(f"   Paragraph {i}: '{text[:30]}...' -> Marker: '{extracted}'")
        
        # Test document processing
        print("\n3. Testing document processing with bullet symbols...")
        
        from core.document_processor import DocumentProcessor
        processor = DocumentProcessor()
        
        with open(test_file, 'rb') as f:
            file_content = f.read()
        
        file_data = {
            'file_content': file_content,
            'filename': 'test_bullet_symbols.docx',
            'tech_stacks': {
                'Python': [
                    'Built scalable backend systems with Django',
                    'Implemented data processing pipelines'
                ]
            }
        }
        
        result = processor.process_document(file_data)
        
        if result.get('success'):
            print("✅ Document processing with bullet symbols succeeded!")
            print(f"   - Points added: {result.get('points_added', 0)}")
            
            # Save and verify the result
            output_file = test_file.replace('.docx', '_processed.docx')
            with open(output_file, 'wb') as f:
                f.write(result['modified_content'])
            
            # Check final consistency
            final_doc = Document(output_file)
            final_marker = formatter.get_document_bullet_marker(final_doc)
            print(f"✅ Final document marker: '{final_marker}'")
            
            # Count bullets
            bullet_count = 0
            for paragraph in final_doc.paragraphs:
                text = paragraph.text.strip()
                if text.startswith('•') or text.startswith('-'):
                    bullet_count += 1
            
            print(f"✅ Total bullet points in final document: {bullet_count}")
            
            return True
        else:
            print(f"❌ Document processing failed: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"❌ Test failed with error: {e}")
        import traceback
        traceback.print_exc()
        return False
    
    finally:
        # Clean up test files
        try:
            if os.path.exists(test_file):
                os.remove(test_file)
            output_file = test_file.replace('.docx', '_processed.docx')
            if os.path.exists(output_file):
                os.remove(output_file)
        except:
            pass

if __name__ == "__main__":
    success = test_bullet_symbol_detection()
    
    if success:
        print("\n🎉 BULLET SYMBOL TEST PASSED! Patch works with bullet symbols.")
    else:
        print("\n❌ BULLET SYMBOL TEST FAILED! Issue with bullet symbol handling.")
    
    exit(0 if success else 1)