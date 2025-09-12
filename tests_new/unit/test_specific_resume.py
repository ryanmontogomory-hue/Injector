"""
Test the emergency patch with the specific resume file that's having issues
"""

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from bullet_consistency_patch import apply_emergency_patch
from formatters.bullet_formatter import BulletFormatter
from core.document_processor import DocumentProcessor
from docx import Document

# Apply the patch
apply_emergency_patch()

def analyze_specific_resume():
    """Analyze the specific resume file that's having issues"""
    
    test_file = r"C:\Users\HP\Downloads\Resume Format 1.docx"
    
    print("=== ANALYZING SPECIFIC RESUME FILE ===")
    print(f"File: {test_file}")
    
    if not os.path.exists(test_file):
        print(f"❌ File not found: {test_file}")
        return False
    
    try:
        # Load the document
        doc = Document(test_file)
        
        print(f"\nDocument has {len(doc.paragraphs)} paragraphs")
        
        # Test the patched formatter
        formatter = BulletFormatter()
        
        print("\n1. DETAILED PARAGRAPH ANALYSIS:")
        bullet_paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                is_bullet = formatter._is_bullet_point(text)
                if is_bullet:
                    marker = formatter._extract_bullet_marker(text)
                    bullet_paragraphs.append((i, text, marker))
                    print(f"   [{i:3d}] BULLET: '{marker}' -> {text[:60]}{'...' if len(text) > 60 else ''}")
                elif len(text) < 50:  # Show short lines (likely headers)
                    print(f"   [{i:3d}] HEADER: {text}")
        
        print(f"\n2. BULLET MARKER DETECTION:")
        detected_marker = formatter.get_document_bullet_marker(doc)
        print(f"   Detected marker: '{detected_marker}'")
        
        # Count each marker type
        marker_counts = {}
        for _, _, marker in bullet_paragraphs:
            marker_counts[marker] = marker_counts.get(marker, 0) + 1
        
        print(f"   Marker distribution:")
        for marker, count in marker_counts.items():
            print(f"     '{marker}': {count} times")
        
        print("\n3. DOCUMENT PROCESSOR TEST:")
        processor = DocumentProcessor()
        
        # Test document-wide detection
        doc_marker = processor._detect_document_bullet_marker(doc)
        print(f"   DocumentProcessor detected: '{doc_marker}'")
        
        print("\n4. FULL PROCESSING TEST:")
        
        with open(test_file, 'rb') as f:
            file_content = f.read()
        
        file_data = {
            'file_content': file_content,
            'filename': 'Resume Format 1.docx',
            'tech_stacks': {
                'TestTech': [
                    'Test bullet point 1 - this should match existing markers',
                    'Test bullet point 2 - this should also match existing markers'
                ]
            }
        }
        
        print("   Processing document with test bullet points...")
        result = processor.process_document(file_data)
        
        if result.get('success'):
            print("   ✅ Document processing succeeded!")
            print(f"   - Points added: {result.get('points_added', 0)}")
            
            # Save and analyze the result
            output_file = test_file.replace('.docx', '_debug_output.docx')
            with open(output_file, 'wb') as f:
                f.write(result['modified_content'])
            
            print(f"   - Output saved: {output_file}")
            
            # Analyze the output
            output_doc = Document(output_file)
            final_marker = formatter.get_document_bullet_marker(output_doc)
            print(f"   - Final marker: '{final_marker}'")
            
            # Check consistency
            print("\n5. CONSISTENCY CHECK:")
            all_bullets = []
            for i, paragraph in enumerate(output_doc.paragraphs):
                text = paragraph.text.strip()
                if formatter._is_bullet_point(text):
                    marker = formatter._extract_bullet_marker(text)
                    all_bullets.append(marker)
            
            unique_markers = set(all_bullets)
            print(f"   Total bullet points: {len(all_bullets)}")
            print(f"   Unique markers used: {unique_markers}")
            
            if len(unique_markers) == 1:
                print("   ✅ CONSISTENCY CHECK PASSED - All bullets use the same marker!")
                return True
            else:
                print("   ❌ CONSISTENCY CHECK FAILED - Multiple markers detected!")
                marker_count = {}
                for marker in all_bullets:
                    marker_count[marker] = marker_count.get(marker, 0) + 1
                for marker, count in marker_count.items():
                    print(f"      '{marker}': {count} times")
                return False
        else:
            print(f"   ❌ Document processing failed: {result.get('error')}")
            return False
            
    except Exception as e:
        print(f"❌ Analysis failed: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = analyze_specific_resume()
    
    print(f"\n=== FINAL RESULT ===")
    if success:
        print("🎉 SPECIFIC RESUME TEST PASSED! Bullet consistency is working.")
    else:
        print("❌ SPECIFIC RESUME TEST FAILED! There are still consistency issues.")
    
    exit(0 if success else 1)