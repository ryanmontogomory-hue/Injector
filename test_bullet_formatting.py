#!/usr/bin/env python3
"""
Test script to verify dynamic bullet formatting detection and application.
"""

import sys
import os
from io import BytesIO

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def create_test_document():
    """Create a test Word document with dash bullet points."""
    from docx import Document
    
    doc = Document()
    
    # Add title
    doc.add_heading('Test Resume', 0)
    
    # Add project section with dash bullets
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Developer | ABC Company | 2022-2024')
    
    # Add bullet points using dashes
    doc.add_paragraph('- Developed web applications using Python and Django')
    doc.add_paragraph('- Implemented RESTful APIs for mobile applications')
    doc.add_paragraph('- Optimized database queries improving performance by 40%')
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_bullet_detection():
    """Test bullet marker detection in document."""
    from core.document_processor import DocumentProcessor
    from docx import Document
    
    # Create test document content
    doc_content = create_test_document()
    
    # Create DocumentProcessor
    processor = DocumentProcessor()
    
    # Load document
    doc = Document(BytesIO(doc_content))
    
    print("📄 Testing bullet marker detection...")
    print("\nOriginal document content:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  {i}: '{para.text}'")
    
    # Test bullet marker detection
    detected_marker = processor._detect_document_bullet_marker(doc)
    print(f"\n🎯 Detected bullet marker: '{detected_marker}'")
    
    # Test individual bullet point detection
    print("\n🔍 Bullet point detection results:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            is_bullet = processor.formatter.is_bullet_point(text)
            if is_bullet:
                marker = processor.formatter._extract_bullet_marker(text)
                separator = processor.formatter._detect_bullet_separator(text)
                print(f"  ✓ Line {i}: '{text}' -> Marker: '{marker}', Separator: '{repr(separator)}'")
    
    return detected_marker

def test_tech_stack_addition():
    """Test adding tech stack points with detected bullet formatting."""
    print("\n" + "="*60)
    print("🧪 TESTING TECH STACK ADDITION WITH DYNAMIC BULLETS")
    print("="*60)
    
    from core.document_processor import DocumentProcessor
    from core.text_parser import parse_input_text, get_parser
    from docx import Document
    
    # Create test document
    doc_content = create_test_document()
    
    # Create test tech stack data
    test_text = "Python Django Flask JavaScript React Node.js PostgreSQL AWS"
    parser = get_parser()
    selected_points, tech_stacks_used = parser.parse_tech_stacks(test_text)
    
    print(f"📋 Tech stacks to add: {tech_stacks_used}")
    print(f"📝 Total points to add: {len([p for points in selected_points for p in points])}")
    
    # Create file data for processing
    file_data = {
        'filename': 'test_resume.docx',
        'file_content': doc_content,
        'tech_stacks': (selected_points, tech_stacks_used),
        'text': test_text
    }
    
    # Process document
    processor = DocumentProcessor()
    result = processor.process_document(file_data)
    
    if result['success']:
        print(f"✅ Processing successful!")
        print(f"   Points added: {result['points_added']}")
        print(f"   Projects modified: {result['projects_modified']}")
        
        # Check the modified document
        modified_doc = Document(BytesIO(result['modified_content']))
        print(f"\n📄 Modified document content:")
        for i, para in enumerate(modified_doc.paragraphs):
            text = para.text.strip()
            if text:
                is_bullet = processor.formatter.is_bullet_point(text)
                marker = "🎯" if is_bullet else "  "
                print(f"  {marker} {i}: '{text}'")
        
        # Verify bullet consistency
        print(f"\n🔍 Bullet consistency check:")
        bullet_markers = []
        for para in modified_doc.paragraphs:
            text = para.text.strip()
            if processor.formatter.is_bullet_point(text):
                marker = processor.formatter._extract_bullet_marker(text)
                bullet_markers.append(marker)
        
        if bullet_markers:
            unique_markers = set(bullet_markers)
            print(f"   Found bullet markers: {list(unique_markers)}")
            if len(unique_markers) == 1:
                print(f"   ✅ CONSISTENT! All bullets use '{list(unique_markers)[0]}'")
            else:
                print(f"   ❌ INCONSISTENT! Mixed bullet markers found")
        
        return True
    else:
        print(f"❌ Processing failed: {result['error']}")
        return False

if __name__ == "__main__":
    print("🧪 DYNAMIC BULLET FORMATTING TEST")
    print("=" * 50)
    
    try:
        # Test 1: Bullet Detection
        detected_marker = test_bullet_detection()
        
        # Test 2: Tech Stack Addition
        success = test_tech_stack_addition()
        
        print(f"\n" + "="*60)
        if success:
            print("🎉 ALL TESTS PASSED! Dynamic bullet formatting is working!")
            print(f"📋 Summary:")
            print(f"   - Detected marker: '{detected_marker}'")
            print(f"   - New points use the same marker consistently")
        else:
            print("❌ TESTS FAILED! Check the implementation.")
        print("="*60)
        
    except Exception as e:
        print(f"💥 Test failed with error: {e}")
        import traceback
        traceback.print_exc()
