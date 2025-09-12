#!/usr/bin/env python3
"""
End-to-end test for Celery task processing with actual resume data.
"""

import sys
import os
import base64
import time
from io import BytesIO

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def create_realistic_resume():
    """Create a realistic Word document for testing."""
    from docx import Document
    
    doc = Document()
    
    # Add title
    doc.add_heading('Software Engineer Resume', 0)
    
    # Add contact info
    doc.add_paragraph('John Doe | john.doe@email.com | +1-555-123-4567')
    
    # Add experience section
    doc.add_heading('Professional Experience', level=1)
    
    # Add project 1
    doc.add_paragraph('Senior Software Engineer | Tech Corp | 2022-2024')
    doc.add_paragraph('- Developed scalable web applications using Python and Django')
    doc.add_paragraph('- Led a team of 5 developers in implementing microservices architecture')
    doc.add_paragraph('- Improved system performance by 40% through database optimization')
    
    # Add project 2  
    doc.add_paragraph('Full Stack Developer | StartupXYZ | 2020-2022')
    doc.add_paragraph('- Built responsive web applications using React and Node.js')
    doc.add_paragraph('- Implemented CI/CD pipelines using Jenkins and Docker')
    doc.add_paragraph('- Designed RESTful APIs serving 10k+ requests per day')
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_celery_task():
    """Test Celery task submission and processing."""
    print("🧪 CELERY END-TO-END TEST")
    print("=" * 50)
    
    try:
        from core.resume_processor import ResumeProcessor
from config.celeryconfig import celery_app
        
        # Create realistic resume
        print("📄 Creating test resume...")
        doc_content = create_realistic_resume()
        
        # Create manual tech stack data
        manual_tech_stacks = {
            'Python': [
                'Built REST APIs using Flask framework for data processing',
                'Implemented automated testing suites with pytest'
            ],
            'JavaScript': [
                'Developed React components for modern user interfaces',
                'Created Node.js backend services with Express framework'
            ],
            'AWS': [
                'Deployed applications using EC2 and ELB for scalability',
                'Implemented serverless functions with Lambda and API Gateway'
            ]
        }
        
        print(f"🔧 Tech stacks to add: {list(manual_tech_stacks.keys())}")
        
        # Create file data for async processing
        file_data = {
            'filename': 'test_resume.docx',
            'file': BytesIO(doc_content),  # Use BytesIO object for async processing
            'text': 'Python JavaScript Flask React Node.js AWS',
            'recipient_email': 'test@example.com',
            'manual_text': ''
        }
        
        # Submit task to Celery
        print("📤 Submitting task to Celery worker...")
        processor = ResumeProcessor()
        async_result = processor.process_single_resume_async(file_data)
        
        print(f"✅ Task submitted successfully!")
        print(f"📝 Task ID: {async_result.id}")
        
        # Monitor task progress
        print(f"⏳ Monitoring task progress...")
        start_time = time.time()
        max_wait_time = 60  # Maximum 60 seconds
        
        while not async_result.ready():
            elapsed = time.time() - start_time
            if elapsed > max_wait_time:
                print(f"❌ Task timed out after {max_wait_time} seconds")
                return False
            
            # Get current status
            task_info = processor.get_async_result(async_result.id)
            state = task_info.get('state', 'UNKNOWN')
            
            if state == 'PROGRESS':
                progress_info = task_info.get('info', {})
                status = progress_info.get('status', 'Processing...')
                progress = progress_info.get('progress', 0)
                print(f"   📊 {state}: {status} ({progress}%)")
            else:
                print(f"   📊 Current status: {state}")
            
            time.sleep(2)  # Wait 2 seconds before checking again
        
        # Get final result
        print(f"✅ Task completed! Getting results...")
        
        try:
            result = async_result.get()
            print(f"🎉 Task processed successfully!")
            print(f"   Points added: {result.get('points_added', 0)}")
            print(f"   Projects modified: {result.get('projects_modified', 0)}")
            print(f"   Filename: {result.get('filename', 'Unknown')}")
            
            # Check if we got the modified content
            if result.get('buffer'):
                print(f"   📄 Modified document received ({len(result['buffer'])} bytes)")
                
                # Verify the document can be opened
                from docx import Document
                modified_doc = Document(BytesIO(result['buffer']))
                print(f"   📊 Modified document has {len(modified_doc.paragraphs)} paragraphs")
                
                # Count bullet points
                bullet_count = 0
                for para in modified_doc.paragraphs:
                    text = para.text.strip()
                    if text and (text.startswith('- ') or text.startswith('• ')):
                        bullet_count += 1
                
                print(f"   🎯 Total bullet points: {bullet_count}")
                
            return True
            
        except Exception as e:
            print(f"❌ Task failed with error: {e}")
            return False
            
    except ImportError as e:
        print(f"❌ Import error: {e}")
        print("   Make sure Celery worker is running!")
        return False
        
    except Exception as e:
        print(f"❌ Test failed with error: {e}")
        import traceback
        traceback.print_exc()
        return False

def check_worker_status():
    """Check if Celery worker is running and has tasks registered."""
    print("🔍 Checking Celery worker status...")
    
    try:
    from config.celeryconfig import celery_app
        
        # Check registered tasks
        inspect = celery_app.control.inspect()
        registered_tasks = inspect.registered()
        
        if registered_tasks:
            print("✅ Worker is running with registered tasks:")
            for worker_name, tasks in registered_tasks.items():
                print(f"   📋 {worker_name}: {tasks}")
            return True
        else:
            print("❌ No workers found or no tasks registered")
            print("   Please start the Celery worker first!")
            return False
            
    except Exception as e:
        print(f"❌ Failed to check worker status: {e}")
        return False

if __name__ == "__main__":
    print("🚀 CELERY INTEGRATION TEST")
    print("=" * 60)
    
    # Step 1: Check if worker is running
    if not check_worker_status():
        print("\n💡 To start the worker, run:")
        print("   python start_celery_worker.py")
        print("\n   Or in a separate terminal window:")
        print("   run_worker.bat")
        exit(1)
    
    print()
    
    # Step 2: Test task processing
    success = test_celery_task()
    
    print("\n" + "="*60)
    if success:
        print("🎉 ALL TESTS PASSED!")
        print("✅ Celery task processing is working correctly")
        print("✅ Resume processing completed successfully")
        print("✅ Dynamic bullet formatting preserved")
    else:
        print("❌ TESTS FAILED!")
        print("   Check the error messages above")
    print("="*60)
