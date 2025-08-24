from docx import Document
from document_processor import DocumentProcessor

# Test with a more realistic resume structure
def test_real_resume_debug():
    print("Testing with realistic resume structure...")
    
    # Create a test document similar to a real resume
    doc = Document()
    doc.add_heading('John Doe - Resume', 0)
    
    # Education section
    doc.add_heading('Education', 1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    doc.add_paragraph('University of Technology, 2015-2019')
    
    # Work Experience section
    doc.add_heading('Work Experience', 1)
    
    # Project 1
    doc.add_paragraph('Company A | Jan 2020 - Present')
    doc.add_paragraph('E-commerce Platform')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Developed web applications using React and Node.js')
    doc.add_paragraph('- Implemented RESTful APIs for product management')
    
    doc.add_paragraph('')  # Empty line
    
    # Project 2
    doc.add_paragraph('Company B | Jun 2018 - Dec 2019')
    doc.add_paragraph('Data Analytics Dashboard')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Built data visualization components with D3.js')
    doc.add_paragraph('- Created backend services with Python and Flask')
    
    doc.add_paragraph('')  # Empty line
    
    # Project 3
    doc.add_paragraph('Company C | Sep 2016 - May 2018')
    doc.add_paragraph('Customer Management System')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Maintained existing codebase in Java')
    doc.add_paragraph('- Wrote unit tests for new features')
    
    print("Document created with 3 projects")
    
    # Create document processor
    processor = DocumentProcessor()
    
    # Detect projects
    projects_data = processor.project_detector.find_projects_and_responsibilities(doc)
    print(f"Detected projects: {projects_data}")
    
    # Convert to structured format
    projects = []
    for i, (title, start_idx, end_idx) in enumerate(projects_data):
        projects.append({
            'title': title,
            'index': i,
            'responsibilities_start': start_idx,
            'responsibilities_end': end_idx
        })
    
    print(f"Structured projects: {projects}")
    
    # Create test tech stacks
    tech_stacks = {
        'React': ['Built responsive UI', 'Used hooks', 'Created components'],
        'MongoDB': ['Created database schema', 'Optimized queries'],
        'Node.js': ['Built REST API', 'Used JWT'],
        'Python': ['Built data processing pipelines', 'Implemented ML models']
    }
    
    print(f"Tech stacks: {tech_stacks}")
    
    # Distribute points
    distribution_result = processor.point_distributor.distribute_points_to_projects(projects, tech_stacks)
    
    print(f"\nDistribution result success: {distribution_result['success']}")
    if not distribution_result['success']:
        print(f"Error: {distribution_result['error']}")
        return
    
    print(f"Points added: {distribution_result['points_added']}")
    print(f"Projects used: {distribution_result['projects_used']}")
    
    print("\nDetailed distribution:")
    for project_title, project_info in distribution_result['distribution'].items():
        print(f"\n{project_title}:")
        print(f"  Project index: {project_info['project_index']}")
        print(f"  Insertion point: {project_info['insertion_point']}")
        print(f"  Responsibilities end: {project_info['responsibilities_end']}")
        print(f"  Total points: {project_info['total_points']}")
        print(f"  Mixed tech stacks: {project_info['mixed_tech_stacks']}")
    
    # Add points to each project
    print("\nAdding points to projects...")
    total_added = 0
    for project_title, project_info in distribution_result['distribution'].items():
        print(f"\nAdding points to {project_title}")
        added = processor.add_points_to_project(doc, project_info)
        print(f"Points added to {project_title}: {added}")
        total_added += added
    
    print(f"\nTotal points added: {total_added}")
    
    # Print document content to see the result
    print("\nDocument content after adding points:")
    for i, para in enumerate(doc.paragraphs):
        print(f"[{i}] {para.text}")

# Run the test
if __name__ == "__main__":
    test_real_resume_debug()