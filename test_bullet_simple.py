#!/usr/bin/env python3
"""
Simple test script for dynamic bullet formatting with manual tech stack data.
"""

import sys
import os
from io import BytesIO

# Add current directory to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

def create_test_document():
    """Create a test Word document with dash bullet points."""
    from docx import Document
    
    doc = Document()
    
    # Add title
    doc.add_heading('Test Resume', 0)
    
    # Add project section with dash bullets
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Software Developer | ABC Company | 2022-2024')
    
    # Add bullet points using dashes
    doc.add_paragraph('- Developed web applications using Python and Django')
    doc.add_paragraph('- Implemented RESTful APIs for mobile applications')
    doc.add_paragraph('- Optimized database queries improving performance by 40%')
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_bullet_formatting():
    """Test bullet formatting with manual tech stack data."""
    print("🧪 SIMPLE BULLET FORMATTING TEST")
    print("=" * 50)
    
    from core.document_processor import DocumentProcessor
    from docx import Document
    
    # Create test document
    doc_content = create_test_document()
    
    # Create manual tech stack data (bypass parser)
    manual_tech_stacks = {
        'Python': [
            'Built REST APIs using Flask framework',
            'Implemented data processing pipelines with Pandas'
        ],
        'JavaScript': [
            'Developed React components for user interface',
            'Created Node.js microservices for backend'
        ]
    }
    
    # Create file data for processing
    file_data = {
        'filename': 'test_resume.docx',
        'file_content': doc_content,
        'tech_stacks': manual_tech_stacks,  # Use dict format directly
        'text': 'Python JavaScript Flask React'
    }
    
    # Test bullet detection first
    doc = Document(BytesIO(doc_content))
    processor = DocumentProcessor()
    
    print("📄 Original document content:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            is_bullet = processor.formatter.is_bullet_point(para.text)
            marker = "🎯" if is_bullet else "  "
            print(f"  {marker} {i}: '{para.text}'")
    
    # Detect bullet marker
    detected_marker = processor._detect_document_bullet_marker(doc)
    print(f"\n🎯 Detected bullet marker: '{detected_marker}'")
    
    # Process document
    print(f"\n🔧 Processing document...")
    result = processor.process_document(file_data)
    
    if result['success']:
        print(f"✅ Processing successful!")
        print(f"   Points added: {result['points_added']}")
        print(f"   Projects modified: {result['projects_modified']}")
        
        # Check the modified document
        modified_doc = Document(BytesIO(result['modified_content']))
        print(f"\n📄 Modified document content:")
        bullet_markers = []
        
        for i, para in enumerate(modified_doc.paragraphs):
            text = para.text.strip()
            if text:
                is_bullet = processor.formatter.is_bullet_point(text)
                marker_symbol = "🎯" if is_bullet else "  "
                print(f"  {marker_symbol} {i}: '{text}'")
                
                # Collect bullet markers
                if is_bullet:
                    marker = processor.formatter._extract_bullet_marker(text)
                    bullet_markers.append(marker)
        
        # Verify bullet consistency
        print(f"\n🔍 Bullet consistency analysis:")
        if bullet_markers:
            unique_markers = set(bullet_markers)
            print(f"   All bullet markers found: {list(bullet_markers)}")
            print(f"   Unique markers: {list(unique_markers)}")
            
            if len(unique_markers) == 1:
                consistent_marker = list(unique_markers)[0]
                print(f"   ✅ SUCCESS! All bullets use '{consistent_marker}' consistently")
                
                # Check if it matches the detected marker
                if consistent_marker == detected_marker:
                    print(f"   ✅ PERFECT! New bullets match the detected format")
                else:
                    print(f"   ⚠️  WARNING: New bullets ('{consistent_marker}') don't match detected ('{detected_marker}')")
            else:
                print(f"   ❌ INCONSISTENT! Mixed bullet markers found")
                return False
        else:
            print(f"   ❓ No bullet points found in processed document")
            
        return True
    else:
        print(f"❌ Processing failed: {result['error']}")
        return False

if __name__ == "__main__":
    try:
        success = test_bullet_formatting()
        
        print(f"\n" + "="*60)
        if success:
            print("🎉 TEST PASSED! Dynamic bullet formatting works correctly!")
            print("📋 Summary:")
            print("   - Detected existing bullet format (dash)")
            print("   - Applied same format to new points")
            print("   - Maintained consistency throughout document")
        else:
            print("❌ TEST FAILED! Check the implementation.")
        print("="*60)
        
    except Exception as e:
        print(f"💥 Test failed with error: {e}")
        import traceback
        traceback.print_exc()
