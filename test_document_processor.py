import os
from docx import Document
from document_processor import DocumentProcessor

# Create a simple test document
def create_test_document():
    doc = Document()
    doc.add_heading('Resume', 0)
    
    # Add Boeing project section
    doc.add_heading('Boeing | Apr 2021 - Present', 1)
    doc.add_paragraph('Software Developer')
    p1 = doc.add_paragraph()
    p1.add_run('Responsibilities:')
    
    # Add some placeholder paragraphs to simulate document structure
    doc.add_paragraph('Placeholder 1')
    doc.add_paragraph('Placeholder 2')
    
    # Add Cardinal Health project section
    doc.add_heading('Cardinal Health | Aug 2017 - Sep 2019', 1)
    doc.add_paragraph('Software Developer')
    p2 = doc.add_paragraph()
    p2.add_run('Responsibilities:')
    
    return doc

# Test the document processor
def test_document_processor():
    print("Testing document processor with project points...")
    
    # Create test document
    doc = create_test_document()
    
    # Create document processor
    processor = DocumentProcessor()
    
    # Create project info for Boeing
    boeing_info = {
        'name': 'Boeing',
        'insertion_point': 4,  # After 'Responsibilities:'
        'responsibilities_end': 10,  # End of responsibilities section
        'mixed_tech_stacks': {
            'React': ['Built responsive UI with React.js', 'Implemented state management with Redux'],
            'Node.js': ['Developed RESTful APIs', 'Implemented authentication with JWT'],
            'AWS': ['Deployed applications to AWS', 'Set up CI/CD pipelines']
        }
    }
    
    # Create project info for Cardinal Health
    cardinal_info = {
        'name': 'Cardinal Health',
        'insertion_point': 20,  # After 'Responsibilities:' in Cardinal Health section
        'responsibilities_end': 25,  # End of responsibilities section
        'mixed_tech_stacks': {
            'React': ['Created reusable React components', 'Used React hooks for state management'],
            'MongoDB': ['Designed MongoDB schemas', 'Optimized database queries'],
            'Python': ['Built data processing pipelines', 'Implemented machine learning models']
        }
    }
    
    # First, add points to Cardinal Health
    print("\nAdding points to Cardinal Health:")
    points_added = processor.add_points_to_project(doc, cardinal_info)
    print(f"Added {points_added} points to Cardinal Health")
    
    # Then add points to Boeing
    print("\nAdding points to Boeing:")
    points_added = processor.add_points_to_project(doc, boeing_info)
    print(f"Added {points_added} points to Boeing")
    
    # Print the document content to verify
    print("\nDocument content after adding points:")
    for i, para in enumerate(doc.paragraphs):
        print(f"[{i}] {para.text}")

# Run the test
if __name__ == "__main__":
    test_document_processor()