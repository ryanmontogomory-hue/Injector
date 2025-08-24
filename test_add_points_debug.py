from docx import Document
from document_processor import DocumentProcessor

# Test adding points to projects with debug output
def test_add_points_debug():
    print("Testing adding points to projects...")
    
    # Create a test document
    doc = Document()
    doc.add_heading('Test Resume', 0)
    
    # Add Project 1
    doc.add_heading('Project 1 | Jan 2020 - Present', 1)
    doc.add_paragraph('Software Developer')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Existing responsibility 1')
    doc.add_paragraph('- Existing responsibility 2')
    
    doc.add_paragraph('')  # Empty line
    
    # Add Project 2
    doc.add_heading('Project 2 | Feb 2018 - Dec 2019', 1)
    doc.add_paragraph('Software Engineer')
    doc.add_paragraph('Responsibilities:')
    doc.add_paragraph('- Existing responsibility A')
    doc.add_paragraph('- Existing responsibility B')
    
    print("Document created with 2 projects")
    
    # Create document processor
    processor = DocumentProcessor()
    
    # Detect projects
    projects_data = processor.project_detector.find_projects_and_responsibilities(doc)
    print(f"Detected projects: {projects_data}")
    
    # Convert to structured format
    projects = []
    for i, (title, start_idx, end_idx) in enumerate(projects_data):
        projects.append({
            'title': title,
            'index': i,
            'responsibilities_start': start_idx,
            'responsibilities_end': end_idx
        })
    
    print(f"Structured projects: {projects}")
    
    # Create test distribution data
    distribution = {
        'Project 1 | Jan 2020 - Present - Software Developer': {
            'mixed_tech_stacks': {
                'React': ['Built responsive UI', 'Used hooks'],
                'Node.js': ['Built REST API']
            },
            'project_index': 0,
            'insertion_point': 4,  # After "Responsibilities:"
            'responsibilities_end': 6,  # After last responsibility
            'total_points': 3
        },
        'Project 2 | Feb 2018 - Dec 2019 - Software Engineer': {
            'mixed_tech_stacks': {
                'MongoDB': ['Created database schema'],
                'Python': ['Built data processing pipelines']
            },
            'project_index': 1,
            'insertion_point': 11,  # After "Responsibilities:"
            'responsibilities_end': 13,  # After last responsibility
            'total_points': 2
        }
    }
    
    print("\nAdding points to projects...")
    
    # Add points to each project
    total_added = 0
    for project_title, project_info in distribution.items():
        print(f"\nAdding points to {project_title}")
        print(f"Project info: {project_info}")
        added = processor.add_points_to_project(doc, project_info)
        print(f"Points added to {project_title}: {added}")
        total_added += added
    
    print(f"\nTotal points added: {total_added}")
    
    # Print document content to see the result
    print("\nDocument content after adding points:")
    for i, para in enumerate(doc.paragraphs):
        print(f"[{i}] {para.text}")

# Run the test
if __name__ == "__main__":
    test_add_points_debug()