"""
Document processing module for Resume Customizer application.
Handles Word document manipulation, project detection, and formatting.
"""

import gc
import threading
from typing import List, Tuple, Dict, Any, Optional
import re
from io import BytesIO
from docx import Document

from config import DOC_CONFIG, PARSING_CONFIG
from infrastructure.utilities.logger import get_logger
from infrastructure.monitoring.performance_cache import cached, cache_key_for_file, get_cache_manager
from infrastructure.monitoring.performance_monitor import performance_decorator
from enhancements.error_handling_enhanced import handle_errors, ErrorSeverity, ErrorHandlerContext
from infrastructure.utilities.memory_optimizer import get_memory_optimizer, with_memory_management
from resume_customizer.formatters.bullet_formatter import BulletFormatter, BulletFormatting

logger = get_logger()


# Precompiled regex patterns for performance (used in tight loops)
CLEAN_NON_ALPHA_RE = re.compile(r'[^a-z ]')
MULTISPACE_RE = re.compile(r'\s+')
NUMBERED_LIST_RE = re.compile(r'^\d+\.')
NUMBERED_LIST_WITH_SPACE_RE = re.compile(r'^\d+\.\s')
UPPER_HEADING_RE = re.compile(r'^[A-Z\s]+$')
MONTHS_RE = re.compile(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', re.IGNORECASE)
DATE_PATTERN_RE = re.compile(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4}|\d{1,2}/\d{1,2}/\d{2,4})\b', re.IGNORECASE)
COMPANY_DATE_PATTERNS = [
    re.compile(r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', re.IGNORECASE),
    re.compile(r'\b\d{4}\b'),
    re.compile(r'\b\d{1,2}/\d{1,2}/\d{2,4}\b'),
    re.compile(r'\b(present|current|now)\b', re.IGNORECASE),
]

class DocumentFormatter:
    """Handles document formatting operations."""
    
    @staticmethod
    @handle_errors("copy_paragraph_formatting", ErrorSeverity.MEDIUM, show_to_user=False)
    def copy_paragraph_formatting(source_para, target_para) -> None:
        """
        Copy all formatting from source paragraph to target paragraph with error handling.
        
        Args:
            source_para: Source paragraph to copy formatting from
            target_para: Target paragraph to apply formatting to
        """
        try:
            # Copy paragraph style
            if source_para.style:
                target_para.style = source_para.style
            
            # Copy paragraph alignment
            if source_para.paragraph_format.alignment is not None:
                target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
            
            # Copy paragraph spacing
            if source_para.paragraph_format.space_before is not None:
                target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
            if source_para.paragraph_format.space_after is not None:
                target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
            if source_para.paragraph_format.line_spacing is not None:
                target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
            
            # Copy indentation
            if source_para.paragraph_format.left_indent is not None:
                target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
            if source_para.paragraph_format.first_line_indent is not None:
                target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        except Exception:
            # Continue if formatting fails - don't let it break the entire process
            pass
    
    @staticmethod
    def copy_run_formatting(source_run, target_run) -> None:
        """
        Copy all formatting from source run to target run.
        
        Args:
            source_run: Source run to copy formatting from
            target_run: Target run to apply formatting to
        """
        try:
            # Copy font properties
            target_run.font.name = source_run.font.name
            target_run.font.size = source_run.font.size
            target_run.font.bold = source_run.font.bold
            target_run.font.italic = source_run.font.italic
            target_run.font.underline = source_run.font.underline
            
            # Copy font color
            if source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb
            
            # Copy highlighting
            if hasattr(source_run.font, 'highlight_color') and source_run.font.highlight_color:
                target_run.font.highlight_color = source_run.font.highlight_color
        except Exception:
            # Continue if formatting fails
            pass



class ProjectDetector:
    """Handles detection of projects and responsibilities sections."""
    
    def __init__(self):
        self.config = PARSING_CONFIG
        
    def find_projects_and_responsibilities(self, doc: Document) -> List[Tuple[str, int, int]]:
        """
        Find all projects and their Responsibilities sections in the resume.
        Supports multiple formats:
        1. Company | Date format with explicit Responsibilities section
        2. Company | Date format with bullet points directly under role
        3. Role titles with bullet points
        
        Args:
            doc: Document to search
            
        Returns:
            List of tuples with (project_name, responsibilities_start_index, responsibilities_end_index)
        """
        projects = []
        current_project = None
        project_title_line = None
        in_responsibilities = False
        responsibilities_start = -1
        found_bullet_points = False
        

        def is_responsibilities_heading(text):
            # Normalize text: lowercase, remove punctuation, collapse spaces
            norm = CLEAN_NON_ALPHA_RE.sub('', text.lower())
            norm = MULTISPACE_RE.sub(' ', norm).strip()
            keywords = [
                'responsibilities', 'key responsibilities', 'duties', 'tasks', 'role', 'position'
            ]
            return any(norm.startswith(k) for k in keywords)

        def is_bullet_point(text):
            """Check if text looks like a bullet point."""
            text = text.strip()
            if not text:
                return False
            # Check for common bullet markers
            bullet_markers = ['•', '-', '*', '◦', '▪', '▫', '‣']
            for marker in bullet_markers:
                if text.startswith(marker + ' ') or text.startswith(marker + '\t') or text.startswith(marker):
                    return True
            # Check for numbered lists
            return NUMBERED_LIST_RE.match(text) is not None

        def is_introductory_paragraph(text):
            """Check if text looks like an introductory paragraph (not a bullet point or heading)."""
            text = text.strip()
            # Skip if it's a bullet point
            if is_bullet_point(text):
                return False
            # Skip if it's a heading (all caps or has specific patterns)
            if text.isupper() or UPPER_HEADING_RE.match(text):
                return False
            # Skip if it's very short (likely a title)
            if len(text.split()) < 5:
                return False
            # Skip if it contains typical heading patterns
            if any(keyword in text.lower() for keyword in ['responsibilities', 'duties', 'role', 'position']):
                return False
            # If it's a longer paragraph, it's likely introductory
            return len(text.split()) >= 10

        def is_project_header(text):
            """Check if text looks like a project/role header."""
            text = text.strip()
            
            # Skip if it's a bullet point
            if is_bullet_point(text):
                return False
                
            # Skip if it's very short (likely just a name)
            if len(text.split()) < 2:
                return False
                
            # Check for company | date format (Kumar S. style)
            if '|' in text and self._looks_like_company_date(text):
                return True
                
            # Check for "Client - Company - Date" format (Viswanadha Raju style)
            if ' - ' in text and MONTHS_RE.search(text):
                return True
                
            # Check for "Role at Company (Location)" format (M. Youssef style)
            if ' at ' in text and '(' in text and ')' in text:
                return True
                
            # Check for role titles with company names
            role_keywords = ['developer', 'engineer', 'manager', 'lead', 'senior', 'software', 'full stack', 'frontend', 'backend']
            if any(keyword in text.lower() for keyword in role_keywords):
                return True
                
            return False

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Look for project headers (multiple formats)
            if is_project_header(text):
                # Save previous project if exists
                if current_project and (responsibilities_start != -1 or found_bullet_points):
                    if responsibilities_start == -1:
                        # If no explicit responsibilities section, use bullet points
                        responsibilities_start = i + 1
                    responsibilities_end = self._find_responsibilities_end(doc, i, responsibilities_start)
                    projects.append((current_project, responsibilities_start, responsibilities_end))
                # Start new project
                current_project = text
                project_title_line = None
                in_responsibilities = False
                responsibilities_start = -1
                found_bullet_points = False

            # If we just found a project header, next non-empty line could be a job title (for Kumar S. format)
            elif current_project and project_title_line is None and text and not is_responsibilities_heading(text) and not is_bullet_point(text) and not is_introductory_paragraph(text):
                project_title_line = text
                # Combine project header with job title for full project name
                current_project = f"{current_project} - {project_title_line}"

            # Look for Responsibilities section (robust match)
            elif text and is_responsibilities_heading(text):
                in_responsibilities = True
                responsibilities_start = i + 1  # Start after the Responsibilities heading

            # Check for introductory paragraph (skip it, bullet points will follow)
            elif current_project and text and is_introductory_paragraph(text):
                # Skip introductory paragraph, continue looking for bullet points
                continue

            # Check for bullet points under current project
            elif current_project and text and is_bullet_point(text):
                found_bullet_points = True
                if responsibilities_start == -1:
                    # Start responsibilities section at first bullet point
                    responsibilities_start = i

            # If we're in responsibilities and find the end (next section or next project)
            elif (in_responsibilities or found_bullet_points) and text and (self._is_section_end(text) or is_project_header(text)):
                if current_project and responsibilities_start != -1:
                    responsibilities_end = i - 1
                    projects.append((current_project, responsibilities_start, responsibilities_end))
                in_responsibilities = False
                responsibilities_start = -1
                found_bullet_points = False

        # Add the last project if found
        if current_project:
            if responsibilities_start != -1 or found_bullet_points:
                if responsibilities_start == -1:
                    # If no explicit responsibilities section, use bullet points
                    responsibilities_start = len(doc.paragraphs) - 1
                responsibilities_end = len(doc.paragraphs) - 1
                projects.append((current_project, responsibilities_start, responsibilities_end))
            else:
                logger.warning(f"No Responsibilities section or bullet points found for project: {current_project}")

        # Defensive filter: remove any projects with responsibilities_start == -1
        projects = [p for p in projects if p[1] != -1]
        
        # Debug logging
        if not projects:
            logger.warning("No projects found. Document structure analysis:")
            for i, para in enumerate(doc.paragraphs[:20]):  # Log first 20 paragraphs
                text = para.text.strip()
                if text:
                    logger.warning(f"Para {i}: '{text[:50]}...' (len: {len(text)})")
        
        return projects
    
    def find_projects(self, doc: Document) -> List[Dict[str, Any]]:
        """Find projects in the document and return structured information."""
        projects_data = self.find_projects_and_responsibilities(doc)

        structured_projects = []
        for i, (title, start_idx, end_idx) in enumerate(projects_data):
            structured_projects.append({
                'title': title,
                'index': i,
                'responsibilities_start': start_idx,
                'responsibilities_end': end_idx
            })

        return structured_projects

    
    def _is_potential_project(self, text: str) -> bool:
        """
        Check if text looks like a project/role heading.
        
        Args:
            text: Text to check
            
        Returns:
            True if text looks like a project heading
        """
        if not text or len(text) >= DOC_CONFIG["max_project_title_length"]:
            return False
            
        if any(keyword in text.lower() for keyword in self.config["project_exclude_keywords"]):
            return False
        
        # Check if it contains pipe symbol (common in project titles like "Project Name | Date")
        if '|' in text:
            return True
        
        # Check if it's all caps (common for project/role headings)
        if text.isupper():
            return True
            
        # Check for common job title keywords
        if any(keyword in text.lower() for keyword in self.config["job_title_keywords"]):
            return True
            
        # Check for project-specific patterns
        if any(keyword in text.lower() for keyword in self.config["project_keywords"]):
            return True
            
        # Check if it looks like a company/project format (contains dates)
        if DATE_PATTERN_RE.search(text):
            return True
            
        return False
    
    def _looks_like_company_date(self, text: str) -> bool:
        """
        Check if text looks like "Company | Date" format.
        
        Args:
            text: Text to check
            
        Returns:
            True if text looks like company|date format
        """
        if not text or '|' not in text:
            return False
        
        parts = text.split('|')
        if len(parts) != 2:
            return False
        
        # Check if second part contains date-like patterns
        date_part = parts[1].strip().lower()
        date_patterns = [
            # kept for readability; use precompiled COMPANY_DATE_PATTERNS
        ]
        
        return any(p.search(date_part) for p in COMPANY_DATE_PATTERNS)
    
    def _is_section_end(self, text: str) -> bool:
        """
        Check if text indicates the end of a section.
        
        Args:
            text: Text to check
            
        Returns:
            True if text indicates section end
        """
        return (
            text.startswith("##") or 
            any(keyword in text.lower() for keyword in self.config["section_end_keywords"])
        )
    
    def _find_responsibilities_end(self, doc: Document, current_index: int, start_index: int) -> int:
        """
        Find the end index of a responsibilities section.
        
        Args:
            doc: Document to search
            current_index: Current paragraph index
            start_index: Start of responsibilities section
            
        Returns:
            End index of responsibilities section
        """
        responsibilities_end = current_index - 1
        for j in range(current_index - 1, start_index - 1, -1):
            if (j < len(doc.paragraphs) and 
                doc.paragraphs[j].text.strip() and 
                not doc.paragraphs[j].text.strip().startswith("-")):
                responsibilities_end = j
                break
        return responsibilities_end


class PointDistributor:
    """Handles distribution of points across projects."""
    
    def distribute_points_to_projects(self, projects: List[Dict], tech_stacks_data) -> Dict[str, Any]:
        """
        Distribute tech stack points across the top 3 projects using round-robin distribution.
        Each project gets a mix of all tech stacks with points distributed evenly.
        This method only calculates distribution, it does NOT add points to the document.
        
        Args:
            projects: List of detected projects
            tech_stacks_data: Either a dictionary of tech stacks or tuple of (points, tech_names)
            
        Returns:
            Dictionary containing distribution results (no points are actually added)
        """
        if not projects or not tech_stacks_data:
            return {'success': False, 'error': 'No projects or tech stacks found'}
        
        # Handle different input formats
        if isinstance(tech_stacks_data, dict):
            tech_stacks = tech_stacks_data
        else:
            # Convert from (points, tech_names) tuple format
            selected_points, tech_names = tech_stacks_data
            if not selected_points or not tech_names:
                return {'success': False, 'error': 'No valid tech stacks found'}
            
            # Create a simple distribution - assign points to tech stacks
            tech_stacks = {}
            points_per_tech = len(selected_points) // len(tech_names) if tech_names else 0
            remaining_points = len(selected_points) % len(tech_names) if tech_names else 0
            
            current_index = 0
            for i, tech_name in enumerate(tech_names):
                points_for_tech = points_per_tech
                if i < remaining_points:
                    points_for_tech += 1
                
                if points_for_tech > 0:
                    tech_stacks[tech_name] = selected_points[current_index:current_index + points_for_tech]
                    current_index += points_for_tech
        
        # Use top 3 projects
        top_projects = projects[:3]
        
        # Calculate round-robin distribution
        distributed_points = self._calculate_round_robin_distribution(tech_stacks, len(top_projects))
        
        distribution = {}
        total_points_added = 0
        
        # Assign distributed points to each project
        for i, project in enumerate(top_projects):
            project_points = distributed_points[i]
            
            distribution[project['title']] = {
                'mixed_tech_stacks': project_points,  # Multiple tech stacks per project
                'project_index': project['index'],
                'insertion_point': project['responsibilities_start'],
                'responsibilities_end': project['responsibilities_end'],
                'total_points': sum(len(points) for points in project_points.values())
            }
            total_points_added += distribution[project['title']]['total_points']
        
        return {
            'success': True,
            'distribution': distribution,
            'points_added': total_points_added,
            'projects_used': len(distribution),
            'distribution_method': 'round_robin'
        }
    
    def _calculate_round_robin_distribution(self, tech_stacks: Dict[str, List[str]], num_projects: int) -> List[Dict[str, List[str]]]:
        """
        Calculate round-robin distribution of tech stack points across projects.
        Each project gets points in a true round-robin fashion, no duplicates, and points are split as evenly as possible.
        """
        project_distributions = [{} for _ in range(num_projects)]
        # Track all points assigned to each project to avoid duplicates
        project_points_set = [set() for _ in range(num_projects)]
        
        # Keep track of the next project to assign points to
        next_project_idx = 0
        
        # Process each tech stack
        for tech_name, points in tech_stacks.items():
            # For each point in this tech stack
            for point in points:
                # Try to find a project that doesn't have this point yet
                start_idx = next_project_idx
                assigned = False
                
                # Try each project once
                for _ in range(num_projects):
                    # If this project doesn't have this point yet
                    if point not in project_points_set[next_project_idx]:
                        # Add the tech stack to this project if it doesn't exist yet
                        if tech_name not in project_distributions[next_project_idx]:
                            project_distributions[next_project_idx][tech_name] = []
                        
                        # Add the point to this project
                        project_distributions[next_project_idx][tech_name].append(point)
                        project_points_set[next_project_idx].add(point)
                        assigned = True
                        
                        # Move to the next project for the next point
                        next_project_idx = (next_project_idx + 1) % num_projects
                        break
                    
                    # Try the next project
                    next_project_idx = (next_project_idx + 1) % num_projects
                
                # If we couldn't assign this point to any project (all projects already have it)
                # Just skip it to avoid duplicates
                if not assigned:
                    continue
        
        return project_distributions


class DocumentProcessor:
    """Main document processor that coordinates all operations."""
    
    def __init__(self):
        self.formatter = BulletFormatter()
        self.project_detector = ProjectDetector()
        self.point_distributor = PointDistributor()
    
    @performance_decorator("project_point_addition", operation_type="doc_modification")
    def add_points_to_project(self, doc: Document, project_info: Dict) -> int:
        """Add tech stack points to a project with proper formatting preservation."""
        if 'insertion_point' not in project_info:
            logger.error(f"project_info missing 'insertion_point': {project_info}")
            return 0
            
        insertion_point = project_info['insertion_point']
        mixed_tech_stacks = project_info['mixed_tech_stacks']
        
        # Find existing bullet formatting in the project section
        formatting_info = self._find_bullet_formatting(doc, project_info)
        
        # Find the best insertion point for new bullets
        insertion_index = self._find_insertion_point(doc, insertion_point)
        
        # Insert all points with proper formatting
        return self._insert_bullet_points(doc, mixed_tech_stacks, insertion_index, formatting_info)
    
    def _find_bullet_formatting(self, doc: Document, project_info: Dict) -> Dict[str, Any]:
        """Find existing bullet formatting in the project section."""
        insertion_point = project_info['insertion_point']
        search_start = max(0, insertion_point - 5)
        search_end = min(len(doc.paragraphs), project_info.get('responsibilities_end', len(doc.paragraphs)) + 5)
        
        # Search for existing bullet formatting in project section
        found_formatting = []
        for i in range(search_start, search_end):
            bullet_formatting = self.formatter.extract_formatting(doc, i)
            if bullet_formatting:
                detected_marker = bullet_formatting.bullet_marker
                if detected_marker:
                    found_formatting.append((detected_marker, bullet_formatting))
        
        # Use the first found marker (prioritize local consistency)
        if found_formatting:
            detected_marker, bullet_formatting = found_formatting[0]
            # Ensure marker consistency (remove extra spaces if any)
            clean_marker = detected_marker.strip() if detected_marker else detected_marker
            logger.info(f"Found local bullet marker in project section: '{clean_marker}'")
            return self._create_formatting_info_from_object(bullet_formatting, clean_marker)
                    
        # Fallback: detect document-wide bullet marker
        marker = self._detect_document_bullet_marker(doc)
        logger.info(f"Using document-wide fallback bullet marker: '{marker}'")
        return self._create_default_formatting(marker)
    
    def _create_formatting_info_from_object(self, bullet_formatting: BulletFormatting, marker: str) -> Dict[str, Any]:
        """Create formatting info from BulletFormatting object."""
        return {
            'formatting': bullet_formatting,
            'fallback_formatting': BulletFormatting(
                runs_formatting=bullet_formatting.runs_formatting,
                paragraph_formatting=bullet_formatting.paragraph_formatting,
                style=bullet_formatting.style,
                bullet_marker=marker,
                bullet_separator=bullet_formatting.bullet_separator,
                list_format=bullet_formatting.list_format
            ),
            'marker': marker
        }
    
    def _create_formatting_info(self, para_formatting: Dict, marker: str) -> Dict[str, Any]:
        """Create formatting info from detected paragraph formatting - legacy support."""
        return {
            'formatting': para_formatting,
            'fallback_formatting': {
                'bullet_marker': marker,
                'bullet_separator': para_formatting.get('bullet_separator', '\t'),
                'style': para_formatting.get('style'),
                'paragraph_formatting': para_formatting.get('paragraph_formatting'),
                'runs_formatting': para_formatting.get('runs_formatting', [])
            },
            'marker': marker
        }
    
    def _create_default_formatting(self, marker: str) -> Dict[str, Any]:
        """Create default formatting when no existing formatting is found."""
        fallback_formatting = BulletFormatting(
            runs_formatting=[],
            paragraph_formatting={},
            style='Normal',
            bullet_marker=marker,
            bullet_separator='\t',
            list_format={
                'ilvl': 0,
                'numId': 1,
                'style': 'List Bullet',
                'indent': 0,
                'is_list': False
            }
        )
        return {
            'formatting': None,
            'fallback_formatting': fallback_formatting,
            'marker': marker
        }
    
    def _find_insertion_point(self, doc: Document, initial_point: int) -> int:
        """Find the best insertion point for new bullet points."""
        # Look for the first bullet point starting from the initial point
        for i in range(initial_point, len(doc.paragraphs)):
            if (doc.paragraphs[i].text.strip() and 
                self.formatter._is_bullet_point(doc.paragraphs[i].text)):
                return i + 1  # Insert after the first bullet point
        
        # If no bullet points found, use the initial point
        return initial_point
    
    def _insert_bullet_points(self, doc: Document, mixed_tech_stacks: Dict, 
                            insertion_index: int, formatting_info: Dict) -> int:
        """Insert bullet points with proper formatting."""
        points_added = 0
        current_index = insertion_index
        
        formatting = formatting_info['formatting']
        fallback_formatting = formatting_info['fallback_formatting']
        marker = formatting_info['marker']
        
        for tech_name, points in mixed_tech_stacks.items():
            for point in points:
                try:
                    # Create new paragraph
                    new_para = self._create_paragraph_at_index(doc, current_index)
                    
                    # Apply formatting
                    self._apply_paragraph_formatting(new_para, formatting, marker)
                    
                    # Add the bullet point content
                    self.formatter.apply_formatting(new_para, formatting, point, fallback_formatting)
                    
                    points_added += 1
                    current_index += 1
                    
                except Exception as e:
                    logger.error(f"Failed to add point '{point}' to project: {str(e)}")
                    continue
        
        return points_added
    
    def _create_paragraph_at_index(self, doc: Document, index: int):
        """Create a new paragraph at the specified index."""
        if index < len(doc.paragraphs):
            return doc.paragraphs[index].insert_paragraph_before()
        else:
            return doc.add_paragraph()
    
    def _apply_paragraph_formatting(self, paragraph, formatting, marker: str) -> None:
        """Apply paragraph-level formatting to a new paragraph."""
        if not formatting:
            return
            
        # Handle both BulletFormatting objects and dictionary formats
        if isinstance(formatting, BulletFormatting):
            # Update bullet marker
            formatting.bullet_marker = marker
            
            # Apply paragraph formatting
            pf_data = formatting.paragraph_formatting
            pf = paragraph.paragraph_format
            for attr, value in pf_data.items():
                if value is not None:
                    try:
                        setattr(pf, attr, value)
                    except Exception:
                        continue
            
            # Apply style
            if formatting.style:
                try:
                    paragraph.style = formatting.style
                except Exception:
                    pass
        else:
            # Legacy dictionary format
            if isinstance(formatting, dict):
                # Update bullet marker
                formatting['bullet_marker'] = marker
                
                # Apply paragraph formatting
                pf_data = formatting.get('paragraph_formatting', {})
                pf = paragraph.paragraph_format
                for attr, value in pf_data.items():
                    if value is not None:
                        try:
                            setattr(pf, attr, value)
                        except Exception:
                            continue
                
                # Apply style
                if formatting.get('style'):
                    try:
                        paragraph.style = formatting['style']
                    except Exception:
                        pass
    
    def _detect_document_bullet_marker(self, doc: Document) -> str:
        """Detect the most common bullet marker used in the entire document with improved accuracy."""
        # Use the improved detection from the bullet formatter
        detected_marker = self.formatter.detect_document_bullet_marker(doc)
        
        # Remove the space from the marker for consistency with existing logic
        if detected_marker.endswith(' '):
            detected_marker = detected_marker.rstrip()
            
        logger.info(f"Detected document bullet marker: '{detected_marker}'")
        return detected_marker
    
    @performance_decorator("document_processing", operation_type="doc_processing")
    def process_document(self, file_data: Dict) -> Dict[str, Any]:
        """
        Process a document by adding tech stack points to projects using round-robin distribution.
        
        Args:
            file_data: Dictionary containing file and processing information
            
        Returns:
            Processing result dictionary
        """
        try:
            # Load document
            doc = Document(BytesIO(file_data['file_content']))
            
            # Detect projects
            projects_data = self.project_detector.find_projects_and_responsibilities(doc)
            if not projects_data:
                return {
                    'success': False,
                    'error': 'No projects found in the document. Please ensure your resume has clear project sections.'
                }
            
            # Convert to structured format
            projects = []
            for i, (title, start_idx, end_idx) in enumerate(projects_data):
                projects.append({
                    'title': title,
                    'index': i,
                    'responsibilities_start': start_idx,
                    'responsibilities_end': end_idx
                })
            
            # Distribute points using round-robin logic
            distribution_result = self.point_distributor.distribute_points_to_projects(projects, file_data['tech_stacks'])
            if not distribution_result['success']:
                return distribution_result
            
            # Add points to each project with mixed tech stacks
            total_added = 0
            # Sort projects by insertion point to process them in order
            sorted_projects = sorted(distribution_result['distribution'].items(),
                                   key=lambda x: x[1]['insertion_point'])
            
            # Keep track of how many paragraphs we've added to adjust insertion points
            paragraph_offset = 0
            
            for project_title, project_info in sorted_projects:
                # Adjust insertion point based on previous additions
                adjusted_project_info = project_info.copy()
                adjusted_project_info['insertion_point'] += paragraph_offset
                if 'responsibilities_end' in adjusted_project_info:
                    adjusted_project_info['responsibilities_end'] += paragraph_offset
                
                added = self.add_points_to_project(doc, adjusted_project_info)
                total_added += added
                
                # Update the offset for subsequent projects
                paragraph_offset += added
            
            # Save modified document
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return {
                'success': True,
                'modified_content': output_buffer.getvalue(),
                'points_added': total_added,
                'projects_modified': distribution_result['projects_used'],
                'distribution_details': distribution_result['distribution'],
                'distribution_method': 'round_robin'
            }
            
        except Exception as e:
            logger.error(f"Document processing failed for {file_data.get('filename', 'unknown')}", exception=e)
            return {
                'success': False,
                'error': f"Failed to process document: {str(e)}"
            }


class FileProcessor:
    """Handles file operations and memory management."""
    
    @staticmethod
    def ensure_file_has_name(file_obj, default_name: str = None) -> Any:
        """
        Ensures file objects have a name attribute for DOCX processing.
        
        Args:
            file_obj: File object to process
            default_name: Default name if none exists
            
        Returns:
            File object with name attribute
        """
        if hasattr(file_obj, 'name'):
            return file_obj
        else:
            # Create wrapper for BytesIO objects
            if default_name is None:
                default_name = DOC_CONFIG["default_filename"]
            file_obj.name = default_name
            return file_obj
    
    @staticmethod
    def cleanup_memory() -> None:
        """Memory optimization - clean up temporary variables."""
        try:
            # Force garbage collection
            gc.collect()
            
            # Clear thread-local storage if any
            current_thread = threading.current_thread()
            if hasattr(current_thread, '__dict__'):
                # Clear non-essential thread data
                thread_dict = current_thread.__dict__.copy()
                for key in thread_dict:
                    if key.startswith('_temp_') or key.startswith('_cache_'):
                        delattr(current_thread, key)
        except Exception:
            # Don't let cleanup errors affect the main application
            pass


def get_document_processor() -> DocumentProcessor:
    """
    Get a new instance of the document processor.
    
    Returns:
        DocumentProcessor instance
    """
    return DocumentProcessor()


def get_streaming_document_processor() -> "StreamingDocumentProcessor":
    """
    Get a new instance of the streaming document processor for large files.
    
    Returns:
        StreamingDocumentProcessor instance optimized for large documents
    """
    return StreamingDocumentProcessor()


def get_optimized_processor(file_size_mb: float = 0):
    """
    Get the best processor based on file size and complexity.
    
    Args:
        file_size_mb: File size in megabytes
        
    Returns:
        Appropriate processor instance (DocumentProcessor or StreamingDocumentProcessor)
    """
    if file_size_mb > 5:  # Files larger than 5MB use streaming processor
        logger.info(f"Using streaming processor for {file_size_mb:.1f}MB file")
        return get_streaming_document_processor()
    else:
        return get_document_processor()


def cleanup_document_resources() -> None:
    """Clean up document processing resources."""
    try:
        # Force garbage collection
        gc.collect()
        
        # Clear any cached document objects
        FileProcessor.cleanup_memory()
        
        # Additional memory cleanup
        import sys
        if hasattr(sys, 'getallocatedblocks'):
            logger.info(f"Memory cleanup completed. Allocated blocks: {sys.getallocatedblocks()}")
        
    except Exception as e:
        # Log but don't fail
        logger.warning(f"Document cleanup failed: {e}")

class StreamingDocumentProcessor:
    """Optimized processor for large documents using streaming techniques."""
    
    def __init__(self, chunk_size: int = 50):
        self.chunk_size = chunk_size
        self.formatter = BulletFormatter()
        self.project_detector = ProjectDetector()
    
    @performance_decorator("streaming_doc_processing", operation_type="stream_processing")
    def process_large_document(self, file_data: Dict) -> Dict[str, Any]:
        """
        Process large documents using streaming approach to minimize memory usage.
        
        Args:
            file_data: Dictionary containing file and processing information
            
        Returns:
            Processing result dictionary
        """
        try:
            # Load document with minimal memory footprint
            doc = Document(BytesIO(file_data['file_content']))
            
            # Process in chunks to avoid memory overload
            total_paragraphs = len(doc.paragraphs)
            
            if total_paragraphs > 200:  # Use streaming for large documents
                logger.info(f"Processing large document with {total_paragraphs} paragraphs using streaming mode")
                return self._process_in_chunks(doc, file_data)
            else:
                # Use regular processing for smaller documents
                regular_processor = DocumentProcessor()
                return regular_processor.process_document(file_data)
        
        except Exception as e:
            logger.error(f"Streaming document processing failed: {e}")
            return {
                'success': False,
                'error': f"Failed to process large document: {str(e)}"
            }
    
    def _process_in_chunks(self, doc: Document, file_data: Dict) -> Dict[str, Any]:
        """Process document in chunks to minimize memory usage."""
        # Find projects first (this is relatively lightweight)
        projects_data = self.project_detector.find_projects_and_responsibilities(doc)
        if not projects_data:
            return {
                'success': False,
                'error': 'No projects found in the large document.'
            }
        
        # Process only the most relevant projects for large documents
        max_projects = min(3, len(projects_data))  # Limit to top 3 projects
        selected_projects = projects_data[:max_projects]
        
        # Convert to structured format
        projects = []
        for i, (title, start_idx, end_idx) in enumerate(selected_projects):
            projects.append({
                'title': title,
                'index': i,
                'responsibilities_start': start_idx,
                'responsibilities_end': end_idx
            })
        
        # Use simplified distribution for large documents
        tech_stacks = file_data.get('tech_stacks', {})
        
        # Distribute points more efficiently for large documents
        simplified_distribution = self._create_simplified_distribution(projects, tech_stacks)
        
        # Apply changes efficiently
        total_added = 0
        for project_info in simplified_distribution:
            try:
                # Add points with memory-efficient approach
                added = self._add_points_efficiently(doc, project_info)
                total_added += added
            except Exception as e:
                logger.warning(f"Failed to add points to project in streaming mode: {e}")
                continue
        
        # Save with memory management
        try:
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            content = output_buffer.getvalue()
            
            # Clean up immediately
            output_buffer.close()
            del doc
            gc.collect()
            
            return {
                'success': True,
                'modified_content': content,
                'points_added': total_added,
                'projects_modified': len(simplified_distribution),
                'distribution_method': 'streaming_optimized',
                'processing_mode': 'streaming'
            }
            
        except Exception as e:
            logger.error(f"Failed to save processed large document: {e}")
            return {
                'success': False,
                'error': f"Failed to save processed document: {str(e)}"
            }
    
    def _create_simplified_distribution(self, projects: List[Dict], tech_stacks: Dict) -> List[Dict]:
        """Create simplified point distribution for large documents."""
        distribution = []
        
        # For large documents, use a simpler distribution strategy
        all_points = []
        for tech_name, points in tech_stacks.items():
            all_points.extend(points[:3])  # Limit points per tech stack
        
        # Distribute evenly across projects
        points_per_project = len(all_points) // len(projects)
        
        for i, project in enumerate(projects):
            start_idx = i * points_per_project
            end_idx = start_idx + points_per_project
            if i == len(projects) - 1:  # Last project gets remaining points
                end_idx = len(all_points)
                
            project_points = all_points[start_idx:end_idx]
            
            if project_points:  # Only include projects with points
                distribution.append({
                    'title': project['title'],
                    'insertion_point': project['responsibilities_start'],
                    'points': project_points
                })
        
        return distribution
    
    def _add_points_efficiently(self, doc: Document, project_info: Dict) -> int:
        """Add points to project with memory-efficient approach but proper formatting."""
        insertion_point = project_info['insertion_point']
        points = project_info['points']
        
        # Find existing bullet formatting around insertion point
        formatting_info = self._find_local_bullet_formatting(doc, insertion_point)
        
        points_added = 0
        for i, point in enumerate(points):
            try:
                # Create paragraph efficiently
                if insertion_point + i < len(doc.paragraphs):
                    new_para = doc.paragraphs[insertion_point + i].insert_paragraph_before()
                else:
                    new_para = doc.add_paragraph()
                
                # Use proper bullet formatting instead of simple text assignment
                self.formatter.apply_formatting(
                    new_para, 
                    formatting_info.get('formatting'),
                    point,
                    formatting_info.get('fallback_formatting')
                )
                
                points_added += 1
                
            except Exception as e:
                logger.warning(f"Failed to add point efficiently: {e}")
                continue
        
        return points_added
    
    def _find_local_bullet_formatting(self, doc: Document, insertion_point: int) -> Dict[str, Any]:
        """Find bullet formatting around the insertion point."""
        # Search around insertion point for existing bullet formatting
        search_start = max(0, insertion_point - 5)
        search_end = min(len(doc.paragraphs), insertion_point + 5)
        
        # Look for existing bullet points with formatting
        for i in range(search_start, search_end):
            if i < len(doc.paragraphs):
                formatting = self.formatter.extract_formatting(doc, i)
                if formatting:
                    logger.info(f"Found local bullet formatting at paragraph {i}: marker='{formatting.bullet_marker}'")
                    return {
                        'formatting': formatting,
                        'fallback_formatting': BulletFormatting(
                            runs_formatting=formatting.runs_formatting,
                            paragraph_formatting=formatting.paragraph_formatting,
                            style=formatting.style,
                            bullet_marker=formatting.bullet_marker,
                            bullet_separator=formatting.bullet_separator,
                            list_format=formatting.list_format
                        )
                    }
        
        # Fallback: detect document-wide marker using formatter
        marker = self._detect_document_bullet_marker_with_formatter(doc)
        logger.info(f"Using document-wide fallback marker: '{marker}'")
        return {
            'formatting': None,
            'fallback_formatting': BulletFormatting(
                runs_formatting=[],
                paragraph_formatting={},
                style='Normal',
                bullet_marker=marker,
                bullet_separator=' ',
                list_format={
                    'ilvl': 0,
                    'numId': 1,
                    'style': 'List Bullet',
                    'indent': 0,
                    'is_list': False
                }
            )
        }
    
    def _detect_document_bullet_marker_with_formatter(self, doc: Document) -> str:
        """Detect bullet marker using the formatter's logic."""
        try:
            marker_counts = {}
            
            # Sample paragraphs throughout the document for better detection
            sample_indices = list(range(0, len(doc.paragraphs), max(1, len(doc.paragraphs) // 20)))
            sample_indices = sample_indices[:20]  # Limit to 20 samples
            
            for i in sample_indices:
                if i < len(doc.paragraphs):
                    para_text = doc.paragraphs[i].text.strip()
                    if para_text and self.formatter._is_bullet_point(para_text):
                        detected_marker = self.formatter._extract_bullet_marker(para_text)
                        if detected_marker and detected_marker != '•':  # Don't count default fallbacks
                            marker_counts[detected_marker] = marker_counts.get(detected_marker, 0) + 1
                        elif detected_marker == '•' and para_text.startswith('•'):
                            # Only count bullet if it's actually in the text
                            marker_counts[detected_marker] = marker_counts.get(detected_marker, 0) + 1
            
            # Return the most common marker
            if marker_counts:
                most_common = max(marker_counts, key=marker_counts.get)
                logger.info(f"Detected document bullet marker: '{most_common}' (count: {marker_counts[most_common]})")
                return most_common
            
            # Default to dash if no markers found (common in resumes)
            logger.info("No bullet markers detected, defaulting to dash (-)")
            return '-'
            
        except Exception as e:
            logger.error(f"Error detecting bullet marker: {e}")
            return '-'
    


def force_memory_cleanup() -> None:
    """Force aggressive memory cleanup when usage is high."""
    try:
        import gc
        import sys
        
        # Multiple garbage collection passes
        for _ in range(3):
            gc.collect()
        
        # Clear any cached objects
        if hasattr(FileProcessor, 'cleanup_memory'):
            FileProcessor.cleanup_memory()
        
        # Clear performance cache if memory is low
        try:
            from infrastructure.monitoring.performance_cache import get_cache_manager
            cache_manager = get_cache_manager()
            
            # Check memory usage before clearing
            try:
                import psutil
                memory = psutil.virtual_memory()
                if memory.percent > 85:  # If memory usage > 85%
                    # Clear oldest cache entries
                    for cache_name in ['document', 'parsing', 'file']:
                        cache = cache_manager.get_cache(cache_name)
                        cache.clear_expired(force=True)  # Clear all expired
                    logger.info("Cache cleared due to high memory usage")
            except ImportError:
                pass
        except Exception:
            pass
        
        logger.info("Aggressive memory cleanup completed")
        
    except Exception as e:
        logger.error(f"Error during aggressive memory cleanup: {e}")


# Global instances and factory functions
_document_processor_instance = None

def get_document_processor() -> DocumentProcessor:
    """Get singleton instance of DocumentProcessor."""
    global _document_processor_instance
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessor()
    return _document_processor_instance

def cleanup_document_resources():
    """Cleanup document processing resources."""
    global _document_processor_instance
    if _document_processor_instance:
        # Clear any cached data
        _document_processor_instance = None
    
    # Force garbage collection
    import gc
    gc.collect()
    logger.info("Document resources cleaned up")
