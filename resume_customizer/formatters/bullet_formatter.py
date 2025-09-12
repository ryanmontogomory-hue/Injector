"""
Bullet point formatter for document processing.
Handles detection and formatting of bullet points in Word documents.
"""
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

from .base_formatters import DocumentFormatter, ListFormatterMixin

@dataclass
class BulletFormatting:
    """Data class to hold bullet formatting information."""
    runs_formatting: List[Dict[str, Any]]
    paragraph_formatting: Dict[str, Any]
    style: str
    bullet_marker: str
    bullet_separator: str
    list_format: Dict[str, Any]


class BulletFormatter(DocumentFormatter, ListFormatterMixin):
    """Handles bullet point formatting and style preservation."""
    
    def __init__(self, bullet_markers: List[str] = None):
        """
        Initialize the BulletFormatter.
        
        Args:
            bullet_markers: List of bullet point markers to recognize
        """
        # Include various dash types commonly used in resumes
        self.dash_variants = ['-', '–', '—', '-']  # hyphen-minus, en-dash, em-dash, hyphen
        self.bullet_markers = bullet_markers or ['•', '●', '◦', '▪', '▫', '‣', '*'] + self.dash_variants
        
        # Common bullet patterns for improved detection
        self.bullet_patterns = [
            r'^\s*[-−–—]\s+',  # Various dash types
            r'^\s*[•·▪▫]\s+',  # Various bullet symbols
            r'^\s*[*]\s+',     # Asterisk
            r'^\s*[+]\s+',     # Plus sign
        ]
        
        # Default marker to use if none found
        self.default_marker = '- '
    
    def extract_formatting(self, doc: DocumentType, paragraph_index: int) -> Optional[BulletFormatting]:
        """
        Extract complete bullet formatting from a paragraph.
        
        Args:
            doc: Document containing the paragraph
            paragraph_index: Index of the paragraph to extract formatting from
            
        Returns:
            BulletFormatting object with formatting information, or None if not a bullet point
        """
        try:
            if paragraph_index >= len(doc.paragraphs):
                return None
                
            para = doc.paragraphs[paragraph_index]
            if not self._is_bullet_point(para.text):
                return None
                
            # Get basic formatting
            formatting_info = BulletFormatting(
                runs_formatting=[],
                paragraph_formatting={},
                style=para.style.name if para.style else 'Normal',
                bullet_marker=self._extract_bullet_marker(para.text) or '•',
                bullet_separator=self._detect_bullet_separator(para.text) or ' ',
                list_format=self._extract_list_format(para)
            )
            
            # Extract formatting from each run
            for run in para.runs:
                try:
                    run_format = {
                        'text': run.text,
                        'font_name': run.font.name if hasattr(run.font, 'name') else None,
                        'font_size': run.font.size.pt if hasattr(run.font, 'size') and run.font.size else None,
                        'bold': run.font.bold if hasattr(run.font, 'bold') else None,
                        'italic': run.font.italic if hasattr(run.font, 'italic') else None,
                        'underline': run.font.underline if hasattr(run.font, 'underline') else None,
                        'color': run.font.color.rgb if (hasattr(run.font, 'color') and 
                                                     run.font.color and 
                                                     hasattr(run.font.color, 'rgb')) else None
                    }
                    formatting_info.runs_formatting.append(run_format)
                except Exception:
                    continue
                    
            # Get paragraph formatting
            if hasattr(para, 'paragraph_format'):
                p_format = para.paragraph_format
                formatting_info.paragraph_formatting = {
                    'alignment': p_format.alignment,
                    'first_line_indent': p_format.first_line_indent,
                    'left_indent': p_format.left_indent,
                    'right_indent': p_format.right_indent,
                    'space_before': p_format.space_before,
                    'space_after': p_format.space_after,
                    'line_spacing': p_format.line_spacing,
                    'keep_together': p_format.keep_together,
                    'keep_with_next': p_format.keep_with_next,
                    'page_break_before': p_format.page_break_before,
                    'widow_control': p_format.widow_control
                }
                
            return formatting_info
            
        except Exception:
            # Try to at least get the bullet marker from the text
            bullet_marker = '-'  # Default to dash instead of bullet symbol
            if 'para' in locals() and para.text:
                try:
                    bullet_marker = self._extract_bullet_marker(para.text)
                    print(f"DEBUG: Exception handler extracted marker: '{bullet_marker}'")
                except Exception:
                    pass
            
            # Return default formatting if anything fails
            return BulletFormatting(
                runs_formatting=[{'text': para.text if 'para' in locals() else ''}],
                paragraph_formatting={},
                style='Normal',
                bullet_marker=bullet_marker,
                bullet_separator=' ',
                list_format={
                    'ilvl': 0,
                    'numId': 1,
                    'style': 'List Bullet',
                    'indent': 0,
                    'is_list': False
                }
            )
    
    def apply_formatting(self, paragraph: Paragraph, formatting: BulletFormatting, 
                        text: str, fallback_formatting: Optional[BulletFormatting] = None) -> None:
        """
        Apply extracted formatting to a new bullet point paragraph.
        
        Args:
            paragraph: Paragraph to format
            formatting: BulletFormatting object with formatting information
            text: Text content to add
            fallback_formatting: Optional fallback formatting if primary formatting fails
        """
        try:
            # Use fallback formatting if main formatting is missing or incomplete
            if not formatting and fallback_formatting:
                formatting = fallback_formatting

            # Apply paragraph style
            if formatting and formatting.style:
                try:
                    paragraph.style = formatting.style
                except Exception:
                    pass

            # Apply paragraph formatting
            if formatting and formatting.paragraph_formatting:
                pf = paragraph.paragraph_format
                for attr, value in formatting.paragraph_formatting.items():
                    if value is not None:
                        try:
                            setattr(pf, attr, value)
                        except Exception:
                            continue

            # Apply Word list formatting if available
            if formatting and formatting.list_format and formatting.list_format.get('is_list'):
                self._apply_list_formatting(paragraph, formatting.list_format)
            
            # Clear existing runs and add formatted text with detected bullet marker
            clean_text = self._clean_bullet_text(text)
            paragraph.clear()
            
            # Get the bullet marker and separator from formatting
            # Ensure we use the detected marker, not the default
            bullet_marker = formatting.bullet_marker if formatting and formatting.bullet_marker else '-'
            bullet_separator = formatting.bullet_separator if formatting and formatting.bullet_separator else ' '
            
            # Validate that we're using the correct marker (preserve detected marker)
            if formatting and formatting.bullet_marker:
                bullet_marker = formatting.bullet_marker  # Explicitly preserve detected marker
            
            # Clean the bullet marker to ensure consistency (remove extra spaces)
            if bullet_marker and bullet_marker != bullet_marker.strip():
                bullet_marker = bullet_marker.strip()
            
            # Add text with proper bullet formatting
            formatted_text = f"{bullet_marker}{bullet_separator}{clean_text}"
            
            # Debug logging to track bullet marker usage
            print(f"DEBUG: apply_formatting using bullet marker: '{bullet_marker}', separator: '{bullet_separator}'")
            print(f"DEBUG: Final formatted text: '{formatted_text[:50]}{'...' if len(formatted_text) > 50 else ''}'")
            
            run = paragraph.add_run(formatted_text)

            # Apply run formatting if available
            if formatting and formatting.runs_formatting:
                # Use the formatting from the first meaningful run
                primary_format = formatting.runs_formatting[0]
                for attr, value in primary_format.items():
                    if attr == 'text' or value is None:
                        continue
                    try:
                        if attr == 'color' and value:
                            run.font.color.rgb = value
                        else:
                            setattr(run.font, attr, value)
                    except Exception:
                        continue

        except Exception as e:
            # Fallback to basic formatting
            print(f"DEBUG: apply_formatting failed, using fallback: {e}")
            self._apply_basic_formatting(paragraph, text, formatting)
    
    def _extract_list_format(self, paragraph: Paragraph) -> Dict[str, Any]:
        """
        Extract list formatting information from a paragraph.
        
        Args:
            paragraph: The paragraph to extract list formatting from
            
        Returns:
            Dictionary containing list formatting information
        """
        list_format = {
            'ilvl': 0,  # Default to top level
            'numId': 1,  # Default to first numbering ID
            'style': 'List Bullet',
            'indent': 0,
            'is_list': False
        }
        
        try:
            # Get style name safely
            if hasattr(paragraph, 'style') and paragraph.style is not None:
                list_format['style'] = paragraph.style.name
            
            # Get indentation
            if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format is not None:
                if hasattr(paragraph.paragraph_format, 'left'):
                    list_format['indent'] = paragraph.paragraph_format.left or 0
            
            # Check if this is actually a bullet point
            if not self._is_bullet_point(paragraph.text):
                return list_format
                
            # If we get here, it's a bullet point
            list_format['is_list'] = True
            
            # Try to get Word list formatting if available
            if not hasattr(paragraph, '_element') or not hasattr(paragraph._element, 'pPr'):
                return list_format
                
            pPr = paragraph._element.pPr
            if pPr is None:
                return list_format
            
            numPr = getattr(pPr, 'numPr', None)
            if numPr is None:
                return list_format
            
            # Get list level (ilvl)
            if hasattr(numPr, 'ilvl') and hasattr(numPr.ilvl, 'val'):
                list_format['ilvl'] = numPr.ilvl.val
            
            # Get numbering ID (numId)
            if hasattr(numPr, 'numId') and hasattr(numPr.numId, 'val'):
                list_format['numId'] = numPr.numId.val
                        
        except Exception:
            pass
            
        return list_format
    
    def _is_bullet_point(self, text: str) -> bool:
        """Check if text starts with a bullet point marker."""
        text = text.strip()
        # Enhanced detection for various bullet formats
        return any(
            text.startswith(marker + ' ') or 
            text.startswith(marker) 
            for marker in self.bullet_markers
        ) or (text and text[0].isdigit() and '.' in text[:3])
    
    def detect_document_bullet_marker(self, document: DocumentType) -> str:
        """
        Detect the primary bullet marker used in the document - improved version
        
        Args:
            document: The Word document to analyze
            
        Returns:
            The detected bullet marker string (e.g., '- ', '• ')
        """
        import re
        
        marker_counts = {}
        
        # Scan all paragraphs for bullet patterns
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                continue
                
            # Check each bullet pattern
            for pattern in self.bullet_patterns:
                match = re.match(pattern, text)
                if match:
                    # Extract the actual marker used
                    marker = match.group().strip() + ' '
                    marker_counts[marker] = marker_counts.get(marker, 0) + 1
                    break
        
        # Return the most common marker, or default if none found
        if marker_counts:
            most_common_marker = max(marker_counts, key=marker_counts.get)
            return most_common_marker
        
        return self.default_marker
    
    def _extract_bullet_marker(self, text: str) -> str:
        """Extract the bullet marker from text - respect document consistency over priority."""
        import re
        
        text = text.strip()
        
        # Use pattern matching for more reliable detection
        for pattern in self.bullet_patterns:
            match = re.match(pattern, text)
            if match:
                # Extract just the marker character, not the spacing
                marker_part = match.group().strip()
                return marker_part if marker_part else '-'
        
        # Check all possible bullet markers without priority bias (include dash variants)
        all_markers = ['•', '●', '◦', '▪', '▫', '‣', '*'] + self.dash_variants
        
        # Find the first marker that matches (respect what's actually in the text)
        for marker in all_markers:
            # Check for marker with separator
            if text.startswith(marker + '\t') or text.startswith(marker + ' '):
                return marker
            # Check for marker without separator (tight formatting)
            elif text.startswith(marker) and len(text) > 1 and not text[1].isalnum():
                return marker
                
        # Check for numbered bullets
        if text and text[0].isdigit():
            for i, char in enumerate(text):
                if char in '.)': 
                    return text[:i+1]
        
        return '•'  # Default to standard bullet only if nothing is found
    
    def _detect_bullet_separator(self, text: str) -> str:
        """Detect whether bullet uses tab or space separator."""
        text = text.strip()
        # Check for common bullet markers including dash variants
        markers = ['•', '●', '◦', '▪', '▫', '‣', '*'] + self.dash_variants
        
        for marker in markers:
            if text.startswith(marker + '\t'):
                return '\t'  # Tab separator
            elif text.startswith(marker + ' '):
                return ' '   # Space separator
        
        return '\t'  # Default to tab for better formatting
    
    def _clean_bullet_text(self, text: str) -> str:
        """Clean bullet text by removing any existing bullet markers."""
        # Remove all possible dash and bullet variants
        dash_and_bullet_chars = '-–—•●*◦▪▫‣ \t'
        clean_text = text.lstrip(dash_and_bullet_chars).lstrip()
        
        # Also handle numbered bullets (e.g., "1. " or "1) ")
        if clean_text and clean_text[0].isdigit():
            for i, char in enumerate(clean_text):
                if char in '.)':
                    clean_text = clean_text[i+1:].lstrip()
                    break
        
        return clean_text
    
    def _apply_basic_formatting(self, paragraph: Paragraph, text: str, 
                              formatting: Optional[BulletFormatting] = None) -> None:
        """Apply basic bullet formatting as a fallback."""
        try:
            # CRITICAL: Ensure we preserve the detected marker even in fallback
            marker = '-'  # Default to dash instead of bullet symbol
            if formatting and formatting.bullet_marker:
                marker = formatting.bullet_marker  # Always use detected marker
                # Clean the marker to ensure no extra spaces
                marker = marker.strip() if marker else '-'
            
            separator = formatting.bullet_separator if formatting and formatting.bullet_separator else ' '
            clean_text = self._clean_bullet_text(text)
            paragraph.clear()
            
            # Create the formatted text with consistent marker
            formatted_text = f"{marker}{separator}{clean_text}"
            paragraph.add_run(formatted_text)
            
            # Log the marker being used for debugging
            print(f"DEBUG: _apply_basic_formatting using marker: '{marker}'")
            
        except Exception as e:
            print(f"DEBUG: _apply_basic_formatting exception: {e}")
            # Last resort: just add the text as-is
            paragraph.clear()
            paragraph.add_run(text)
