"""
Memory optimization utilities for Resume Customizer application.
Provides automatic memory cleanup, monitoring, and resource management.
"""

import gc
import weakref
import threading
import time
from contextlib import contextmanager
from typing import Dict, List, Optional, Any, Callable
from dataclasses import dataclass, field
from datetime import datetime, timedelta

try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False

from logger import get_logger

logger = get_logger()


@dataclass
class MemoryStats:
    """Memory usage statistics."""
    used_mb: float
    available_mb: float
    percent_used: float
    process_memory_mb: float
    timestamp: datetime


class ResourceTracker:
    """Track and manage application resources."""
    
    def __init__(self):
        self._active_resources = weakref.WeakSet()
        self._resource_types = {}
        self._lock = threading.Lock()
    
    def register_resource(self, resource: Any, resource_type: str = "unknown"):
        """Register a resource for tracking."""
        with self._lock:
            self._active_resources.add(resource)
            self._resource_types[id(resource)] = resource_type
    
    def get_active_count(self, resource_type: Optional[str] = None) -> int:
        """Get count of active resources."""
        if resource_type is None:
            return len(self._active_resources)
        
        count = 0
        for resource in self._active_resources:
            if self._resource_types.get(id(resource)) == resource_type:
                count += 1
        return count
    
    def cleanup_resources(self, resource_type: Optional[str] = None):
        """Force cleanup of resources."""
        resources_to_cleanup = []
        
        for resource in list(self._active_resources):
            if resource_type is None or self._resource_types.get(id(resource)) == resource_type:
                resources_to_cleanup.append(resource)
        
        # Try to cleanup each resource
        for resource in resources_to_cleanup:
            try:
                if hasattr(resource, 'close'):
                    resource.close()
                elif hasattr(resource, '__exit__'):
                    resource.__exit__(None, None, None)
            except Exception as e:
                logger.warning(f"Failed to cleanup resource {type(resource)}: {e}")


class MemoryMonitor:
    """Monitor memory usage and provide alerts."""
    
    def __init__(self, alert_threshold: float = 85.0):
        self.alert_threshold = alert_threshold
        self._last_alert = None
        self._alert_cooldown = timedelta(minutes=5)
    
    def get_memory_stats(self) -> MemoryStats:
        """Get current memory statistics."""
        if PSUTIL_AVAILABLE:
            memory = psutil.virtual_memory()
            process = psutil.Process()
            process_memory = process.memory_info().rss / (1024 * 1024)
            
            return MemoryStats(
                used_mb=memory.used / (1024 * 1024),
                available_mb=memory.available / (1024 * 1024),
                percent_used=memory.percent,
                process_memory_mb=process_memory,
                timestamp=datetime.now()
            )
        else:
            # Fallback when psutil is not available
            return MemoryStats(
                used_mb=0.0,
                available_mb=0.0,
                percent_used=0.0,
                process_memory_mb=0.0,
                timestamp=datetime.now()
            )
    
    def check_memory_pressure(self) -> bool:
        """Check if system is under memory pressure."""
        stats = self.get_memory_stats()
        
        if stats.percent_used > self.alert_threshold:
            now = datetime.now()
            if (self._last_alert is None or 
                now - self._last_alert > self._alert_cooldown):
                
                logger.warning(
                    f"High memory usage detected: {stats.percent_used:.1f}% "
                    f"({stats.used_mb:.0f}MB used, {stats.available_mb:.0f}MB available)"
                )
                self._last_alert = now
                return True
        
        return False
    
    def suggest_cleanup(self) -> List[str]:
        """Suggest cleanup actions based on memory usage."""
        stats = self.get_memory_stats()
        suggestions = []
        
        if stats.percent_used > 90:
            suggestions.extend([
                "Critical: Restart the application immediately",
                "Close other browser tabs and applications",
                "Process fewer files at once"
            ])
        elif stats.percent_used > 75:
            suggestions.extend([
                "Process files in smaller batches",
                "Clear application cache",
                "Close unused browser tabs"
            ])
        elif stats.percent_used > 60:
            suggestions.append("Consider processing fewer files simultaneously")
        
        return suggestions


class DocumentResourceManager:
    """Manage Word document resources with automatic cleanup."""
    
    def __init__(self):
        self._open_documents = {}
        self._resource_tracker = ResourceTracker()
        self._lock = threading.Lock()
    
    @contextmanager
    def managed_document(self, file_path: str = None, file_obj: Any = None):
        """Context manager for Word documents with automatic cleanup."""
        from docx import Document
        
        doc = None
        doc_id = None
        
        try:
            # Create document
            if file_obj:
                doc = Document(file_obj)
            elif file_path:
                doc = Document(file_path)
            else:
                doc = Document()
            
            # Track the document
            doc_id = id(doc)
            with self._lock:
                self._open_documents[doc_id] = {
                    'document': doc,
                    'created_at': datetime.now(),
                    'file_path': file_path
                }
            
            self._resource_tracker.register_resource(doc, "word_document")
            
            yield doc
            
        finally:
            # Cleanup
            if doc_id and doc_id in self._open_documents:
                with self._lock:
                    del self._open_documents[doc_id]
            
            # Force garbage collection for document
            if doc:
                del doc
            gc.collect()
    
    def cleanup_all_documents(self):
        """Force cleanup of all tracked documents."""
        with self._lock:
            doc_count = len(self._open_documents)
            self._open_documents.clear()
        
        self._resource_tracker.cleanup_resources("word_document")
        gc.collect()
        
        logger.info(f"Cleaned up {doc_count} Word documents")
    
    def get_open_document_count(self) -> int:
        """Get count of currently open documents."""
        return len(self._open_documents)


class MemoryOptimizer:
    """Main memory optimization coordinator."""
    
    def __init__(self):
        self.monitor = MemoryMonitor()
        self.resource_tracker = ResourceTracker()
        self.doc_manager = DocumentResourceManager()
        self._cleanup_interval = 300  # 5 minutes
        self._last_cleanup = None
    
    def optimize_memory(self, force: bool = False) -> Dict[str, Any]:
        """Perform memory optimization."""
        now = datetime.now()
        
        # Check if it's time for cleanup
        if (not force and self._last_cleanup and 
            now - self._last_cleanup < timedelta(seconds=self._cleanup_interval)):
            return {"status": "skipped", "reason": "cleanup_interval_not_reached"}
        
        logger.info("Starting memory optimization...")
        
        # Get before stats
        before_stats = self.monitor.get_memory_stats()
        
        # Cleanup steps
        cleanup_results = {}
        
        # 1. Force garbage collection
        collected = gc.collect()
        cleanup_results['garbage_collected'] = collected
        
        # 2. Cleanup documents
        doc_count_before = self.doc_manager.get_open_document_count()
        self.doc_manager.cleanup_all_documents()
        cleanup_results['documents_cleaned'] = doc_count_before
        
        # 3. Clear weak references
        self.resource_tracker.cleanup_resources()
        
        # 4. Additional Python optimizations
        if hasattr(gc, 'set_threshold'):
            gc.set_threshold(700, 10, 10)  # More aggressive GC
        
        # Get after stats
        after_stats = self.monitor.get_memory_stats()
        
        # Calculate savings
        memory_saved = before_stats.process_memory_mb - after_stats.process_memory_mb
        
        self._last_cleanup = now
        
        result = {
            "status": "completed",
            "timestamp": now,
            "before_memory_mb": before_stats.process_memory_mb,
            "after_memory_mb": after_stats.process_memory_mb,
            "memory_saved_mb": memory_saved,
            "cleanup_actions": cleanup_results
        }
        
        logger.info(
            f"Memory optimization completed. "
            f"Saved {memory_saved:.1f}MB "
            f"({cleanup_results.get('garbage_collected', 0)} objects collected)"
        )
        
        return result
    
    def auto_cleanup_if_needed(self):
        """Automatically cleanup if memory pressure is detected."""
        if self.monitor.check_memory_pressure():
            return self.optimize_memory(force=True)
        return None
    
    def get_optimization_report(self) -> Dict[str, Any]:
        """Get comprehensive memory optimization report."""
        stats = self.monitor.get_memory_stats()
        suggestions = self.monitor.suggest_cleanup()
        
        return {
            "memory_stats": stats,
            "open_documents": self.doc_manager.get_open_document_count(),
            "active_resources": self.resource_tracker.get_active_count(),
            "suggestions": suggestions,
            "last_cleanup": self._last_cleanup,
            "needs_cleanup": stats.percent_used > self.monitor.alert_threshold
        }


# Global instance
_memory_optimizer = None

def get_memory_optimizer() -> MemoryOptimizer:
    """Get singleton memory optimizer instance."""
    global _memory_optimizer
    if _memory_optimizer is None:
        _memory_optimizer = MemoryOptimizer()
    return _memory_optimizer


# Decorator for automatic memory management
def with_memory_management(cleanup_after: bool = True):
    """Decorator to add automatic memory management to functions."""
    def decorator(func):
        def wrapper(*args, **kwargs):
            optimizer = get_memory_optimizer()
            
            # Check memory before execution
            optimizer.auto_cleanup_if_needed()
            
            try:
                result = func(*args, **kwargs)
                return result
            finally:
                if cleanup_after:
                    # Force cleanup after processing
                    optimizer.optimize_memory(force=False)
        
        return wrapper
    return decorator


# Context manager for bulk operations
@contextmanager
def memory_efficient_batch(batch_size: int = 5):
    """Context manager for memory-efficient batch processing."""
    optimizer = get_memory_optimizer()
    
    try:
        # Initial cleanup
        optimizer.auto_cleanup_if_needed()
        
        yield batch_size
        
    finally:
        # Final cleanup
        optimizer.optimize_memory(force=True)
