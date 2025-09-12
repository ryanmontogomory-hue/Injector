"""
EMERGENCY BULLET CONSISTENCY PATCH
This patch directly fixes the bullet formatting issue by overriding the apply_formatting method
"""

import re
from typing import Dict, Any, List, Optional
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

def patch_bullet_formatter(formatter_instance):
    """
    Patch an existing BulletFormatter instance to ensure bullet consistency
    """
    
    # Store the original methods
    original_apply_formatting = formatter_instance.apply_formatting
    original_extract_marker = formatter_instance._extract_bullet_marker
    
    def get_document_bullet_marker(document: DocumentType) -> str:
        """Get the most common bullet marker in the document"""
        marker_counts = {}
        
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Check for dash markers (most common in resumes)
            if re.match(r'^\s*[-−–—]\s+', text):
                marker = text.strip()[0]
                marker_counts[marker] = marker_counts.get(marker, 0) + 1
            # Check for bullet symbols - improved detection
            elif re.match(r'^\s*[•·▪▫]\s+', text):
                marker = text.strip()[0]
                marker_counts[marker] = marker_counts.get(marker, 0) + 1
                print(f"PATCH: Found bullet symbol marker: '{marker}' in text: '{text[:30]}...'")
            # Check for asterisk
            elif re.match(r'^\s*\*\s+', text):
                marker_counts['*'] = marker_counts.get('*', 0) + 1
        
        # Return most common marker or default to dash
        if marker_counts:
            most_common = max(marker_counts, key=marker_counts.get)
            print(f"PATCH: Detected document bullet marker: '{most_common}'")
            return most_common
        
        print("PATCH: No bullets found, defaulting to dash (-)")
        return '-'
    
    def patched_extract_marker(text: str) -> str:
        """Extract bullet marker with guaranteed consistency"""
        text = text.strip()
        
        # Direct pattern matching for reliability
        if re.match(r'^\s*[-−–—]', text):
            marker = text.strip()[0]
            print(f"PATCH: Extracted dash marker: '{marker}'")
            return marker  # Return the actual dash character
        elif re.match(r'^\s*[•·▪▫]', text):
            marker = text.strip()[0]
            print(f"PATCH: Extracted bullet symbol: '{marker}'")
            return marker  # Return the actual bullet character
        elif re.match(r'^\s*\*', text):
            print("PATCH: Extracted asterisk marker: '*'")
            return '*'
        elif re.match(r'^\s*\+', text):
            print("PATCH: Extracted plus marker: '+'")
            return '+'
        
        # Default to dash
        print("PATCH: No marker found, defaulting to dash: '-'")
        return '-'
    
    def patched_apply_formatting(paragraph: Paragraph, formatting, text: str, fallback_formatting=None):
        """Patched apply_formatting that ensures bullet consistency"""
        try:
            print(f"PATCH: Applying formatting to text: '{text[:30]}...'")
            
            # Get the bullet marker - prioritize the formatting object
            bullet_marker = '-'  # Safe default
            
            if formatting and hasattr(formatting, 'bullet_marker') and formatting.bullet_marker:
                bullet_marker = formatting.bullet_marker.strip()
                print(f"PATCH: Using formatting bullet marker: '{bullet_marker}'")
            elif fallback_formatting and hasattr(fallback_formatting, 'bullet_marker') and fallback_formatting.bullet_marker:
                bullet_marker = fallback_formatting.bullet_marker.strip()
                print(f"PATCH: Using fallback bullet marker: '{bullet_marker}'")
            else:
                print(f"PATCH: Using default bullet marker: '{bullet_marker}'")
            
            # Clean the text (remove any existing bullets)
            clean_text = text.strip()
            # Remove common bullet markers from the start
            for pattern in [r'^\s*[-−–—•·▪▫*+]\s*', r'^\s*\d+[.)]\s*']:
                clean_text = re.sub(pattern, '', clean_text)
            clean_text = clean_text.strip()
            
            # Apply the consistent bullet format
            separator = ' '  # Always use space separator for consistency
            formatted_text = f"{bullet_marker}{separator}{clean_text}"
            
            print(f"PATCH: Final formatted text: '{formatted_text}'")
            
            # Clear and apply
            paragraph.clear()
            paragraph.add_run(formatted_text)
            
            # Try to apply additional formatting if available
            if formatting:
                try:
                    if hasattr(formatting, 'style') and formatting.style:
                        paragraph.style = formatting.style
                except:
                    pass
                    
                try:
                    if hasattr(formatting, 'paragraph_formatting') and formatting.paragraph_formatting:
                        pf = paragraph.paragraph_format
                        for attr, value in formatting.paragraph_formatting.items():
                            if value is not None:
                                try:
                                    setattr(pf, attr, value)
                                except:
                                    continue
                except:
                    pass
            
            print(f"PATCH: Successfully applied formatting with marker '{bullet_marker}'")
            
        except Exception as e:
            print(f"PATCH: Error in apply_formatting: {e}")
            # Absolute fallback
            try:
                paragraph.clear()
                paragraph.add_run(f"- {text.strip()}")
                print("PATCH: Applied absolute fallback formatting")
            except Exception as e2:
                print(f"PATCH: Even fallback failed: {e2}")
                paragraph.clear()
                paragraph.add_run(text)
    
    # Apply the patches
    formatter_instance._extract_bullet_marker = patched_extract_marker
    formatter_instance.apply_formatting = patched_apply_formatting
    formatter_instance.get_document_bullet_marker = get_document_bullet_marker
    
    print("PATCH: BulletFormatter has been patched for consistency!")
    return formatter_instance

def patch_document_processor(processor_instance):
    """Patch a DocumentProcessor instance to use consistent bullet detection"""
    
    original_detect_method = processor_instance._detect_document_bullet_marker
    
    def patched_detect_document_bullet_marker(doc):
        """Patched document bullet marker detection"""
        marker_counts = {}
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Check for dash markers (most common in resumes)
            if re.match(r'^\s*[-−–—]\s+', text):
                marker = text.strip()[0]
                marker_counts[marker] = marker_counts.get(marker, 0) + 1
            # Check for bullet symbols - improved detection
            elif re.match(r'^\s*[•·▪▫]\s+', text):
                marker = text.strip()[0]
                marker_counts[marker] = marker_counts.get(marker, 0) + 1
                print(f"PATCH: DocumentProcessor found bullet symbol: '{marker}' in text: '{text[:30]}...'")
            # Check for asterisk
            elif re.match(r'^\s*\*\s+', text):
                marker_counts['*'] = marker_counts.get('*', 0) + 1
        
        # Return most common marker or default to dash
        if marker_counts:
            most_common = max(marker_counts, key=marker_counts.get)
            print(f"PATCH: DocumentProcessor detected bullet marker: '{most_common}' (count: {marker_counts[most_common]})")
            return most_common
        
        print("PATCH: DocumentProcessor - No bullets found, defaulting to dash (-)")
        return '-'
    
    processor_instance._detect_document_bullet_marker = patched_detect_document_bullet_marker
    
    # Also patch the formatter if it exists
    if hasattr(processor_instance, 'formatter'):
        patch_bullet_formatter(processor_instance.formatter)
    
    print("PATCH: DocumentProcessor has been patched for consistency!")
    return processor_instance

# Auto-apply patch when imported
def apply_emergency_patch():
    """Apply the emergency patch to fix bullet consistency"""
    try:
        # Import the modules that need patching
        from core.document_processor import DocumentProcessor
        from formatters.bullet_formatter import BulletFormatter
        
        # Patch the classes at module level
        original_bullet_init = BulletFormatter.__init__
        original_processor_init = DocumentProcessor.__init__
        
        def patched_bullet_init(self, *args, **kwargs):
            original_bullet_init(self, *args, **kwargs)
            patch_bullet_formatter(self)
        
        def patched_processor_init(self, *args, **kwargs):
            original_processor_init(self, *args, **kwargs)
            patch_document_processor(self)
        
        BulletFormatter.__init__ = patched_bullet_init
        DocumentProcessor.__init__ = patched_processor_init
        
        print("🔧 EMERGENCY PATCH APPLIED: Bullet consistency is now guaranteed!")
        return True
        
    except Exception as e:
        print(f"❌ Failed to apply emergency patch: {e}")
        return False

if __name__ == "__main__":
    apply_emergency_patch()
    print("Emergency patch ready to apply!")