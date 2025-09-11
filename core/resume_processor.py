"""
Resume processing module for Resume Customizer application.
Coordinates all resume processing operations with improved architecture and error handling.
"""

import threading
import functools
import queue
import os
import time
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional, Callable, Tuple, Union
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass, field

# External imports
from docx import Document
import streamlit as st

# Local imports
from audit_logger import audit_logger
from utilities.logger import get_logger
from core.text_parser import parse_input_text, get_parser
from core.document_processor import get_document_processor, FileProcessor
from core.email_handler import get_email_manager
from config import ERROR_MESSAGES, SUCCESS_MESSAGES, APP_CONFIG
from security_enhancements import rate_limit
from utilities.structured_logger import get_structured_logger, with_structured_logging, processing_logger
from enhancements.error_handling_enhanced import handle_errors, ErrorSeverity, ErrorHandlerContext
from monitoring.circuit_breaker import file_processing_circuit_breaker
from enhancements.enhanced_error_recovery import RobustResumeProcessor, get_error_recovery_manager
from monitoring.distributed_cache import cached_processing, get_distributed_cache_manager
from enhancements.metrics_analytics_enhanced import record_resume_processed, record_metrics
from enhancements.progress_tracker_enhanced import get_progress_tracker

logger = get_logger()

# Constants
CACHE_TTL_SECONDS = 300  # 5 minutes
MAX_CACHE_SIZE = 100
DEFAULT_PROCESSING_TIMEOUT = 30


@dataclass
class ProcessingResult:
    """Data class for processing results with comprehensive information."""
    success: bool
    filename: str
    error: Optional[str] = None
    points_added: int = 0
    processing_time: float = 0.0
    tech_stacks_used: List[str] = field(default_factory=list)
    modified_content: Optional[bytes] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class CacheEntry:
    """Cache entry with TTL support."""
    result: ProcessingResult
    timestamp: datetime
    access_count: int = 0
    
    def is_expired(self) -> bool:
        """Check if cache entry is expired."""
        return datetime.now() - self.timestamp > timedelta(seconds=CACHE_TTL_SECONDS)


def profile_step(step_name: str):
    """Decorator to profile and log slow operations."""
    def decorator(func):
        @functools.wraps(func)
        def wrapper(*args, **kwargs):
            start = time.time()
            try:
                result = func(*args, **kwargs)
                duration = time.time() - start
                if duration > 1.0:  # Log only slow steps (>1s)
                    logger.info(f"[PROFILE] {step_name} took {duration:.2f}s")
                return result
            except Exception as e:
                duration = time.time() - start
                logger.error(f"[PROFILE] {step_name} failed after {duration:.2f}s: {e}")
                raise
        return wrapper
    return decorator



def _process_single_resume_worker(payload: Dict[str, Any]) -> ProcessingResult:
    """Worker function for processing a single resume in a separate process.
    
    Args:
        payload: Serializable dictionary containing file data and processing parameters
        
    Returns:
        ProcessingResult with processing outcome
    """
    start_time = time.time()
    filename = payload.get('filename', 'unknown')
    
    try:
        # Validate payload
        if 'file_content' not in payload:
            return ProcessingResult(
                success=False,
                filename=filename,
                error="Missing file content in payload",
                processing_time=time.time() - start_time
            )
        
        # Reconstruct file object from bytes
        file_obj = BytesIO(payload['file_content'])
        
        # Create processor and process resume
        processor = ResumeProcessor()
        file_data = {
            'filename': filename,
            'file': file_obj,
            'text': payload.get('text', ''),
            'recipient_email': payload.get('recipient_email', ''),
            'sender_email': payload.get('sender_email', ''),
            'sender_password': payload.get('sender_password', ''),
            'smtp_server': payload.get('smtp_server', ''),
            'smtp_port': payload.get('smtp_port', 465),
            'email_subject': payload.get('email_subject', ''),
            'email_body': payload.get('email_body', ''),
        }
        
        result = processor.process_single_resume(file_data)
        
        # Convert to ProcessingResult if it's not already
        if isinstance(result, dict):
            return ProcessingResult(
                success=result.get('success', False),
                filename=filename,
                error=result.get('error'),
                points_added=result.get('points_added', 0),
                processing_time=time.time() - start_time,
                tech_stacks_used=result.get('tech_stacks_used', []),
                modified_content=result.get('modified_content'),
                metadata=result.get('metadata', {})
            )
        return result
        
    except Exception as e:
        logger.error(f"Worker error processing {filename}: {e}")
        return ProcessingResult(
            success=False,
            filename=filename,
            error=str(e),
            processing_time=time.time() - start_time
        )


class ResumeProcessor:
    def process_single_resume_async(self, file_data: Dict[str, Any]):
        """
        Submit a resume processing job to Celery and return the AsyncResult.
        Serializes file content to base64 for JSON compatibility.
        """
        try:
            # Use Celery app to get the registered task instead of direct import
            # This avoids circular import issues
            from celeryconfig import celery_app
            process_resume_task = celery_app.tasks.get('tasks.process_resume_task')
            
            if process_resume_task is None:
                # Try to import the tasks module to register tasks
                try:
                    import tasks  # This should register the task
                    process_resume_task = celery_app.tasks.get('tasks.process_resume_task')
                except ImportError:
                    pass
            
            if process_resume_task is None:
                raise ImportError("Task 'tasks.process_resume_task' not found in Celery app")
                
        except ImportError as import_error:
            import os
            import sys
            cwd = os.getcwd()
            python_path = sys.path[:3]  # Show first 3 paths
            error_msg = f"""Celery tasks module not found. Debugging info:
            - Current working directory: {cwd}
            - Python path (first 3): {python_path}
            - Import error: {str(import_error)}
            - Make sure tasks.py exists and Celery worker is running
            - Try running from project root directory"""
            raise RuntimeError(error_msg)
        
        try:
            # Create a copy of file_data to avoid modifying the original
            serialized_data = file_data.copy()
            
            # Serialize the file object to base64 for JSON compatibility
            if 'file' in serialized_data and serialized_data['file'] is not None:
                file_obj = serialized_data['file']
                
                # Read file content
                file_obj.seek(0)  # Ensure we're at the start of the file
                file_content = file_obj.read()
                file_obj.seek(0)  # Reset position for potential future use
                
                # Convert to base64 string
                import base64
                file_content_b64 = base64.b64encode(file_content).decode('utf-8')
                
                # Replace file object with base64 content
                serialized_data['file_content_b64'] = file_content_b64
                del serialized_data['file']  # Remove non-serializable file object
            else:
                raise ValueError("File object is required for async processing")
            
            return process_resume_task.delay(serialized_data)
            
        except Exception as e:
            error_msg = str(e)
            if 'Redis' in error_msg or 'Connection refused' in error_msg:
                raise RuntimeError("Celery broker (Redis) connection failed. Make sure Redis is running or configure an alternative broker.")
            elif 'NoneType' in error_msg and 'Redis' in error_msg:
                raise RuntimeError("Redis configuration issue. Check celeryconfig.py and ensure Redis is properly installed.")
            elif 'not JSON serializable' in error_msg:
                raise RuntimeError("File serialization failed. Please try again or use synchronous processing.")
            else:
                raise RuntimeError(f"Celery task submission failed: {error_msg}")

    def get_async_result(self, task_id: str):
        """
        Get the result/status of a Celery async task by task_id.
        """
        try:
            from celeryconfig import celery_app
        except ImportError:
            raise RuntimeError("Celery config not found. Make sure celeryconfig.py exists and Celery is installed.")
        
        try:
            async_result = celery_app.AsyncResult(task_id)
            
            # Get the state safely
            state = async_result.state
            
            # Get result safely, handling different states
            result = None
            if async_result.ready():
                try:
                    result = async_result.result
                except Exception as result_error:
                    # If result retrieval fails, include error info
                    result = {
                        'error': f"Result retrieval failed: {str(result_error)}",
                        'success': False
                    }
            
            # Get additional task info if available
            task_info = {
                'state': state,
                'result': result
            }
            
            # Add progress info if available (for PROGRESS state)
            if state == 'PROGRESS' and hasattr(async_result, 'info') and async_result.info:
                task_info['info'] = async_result.info
            
            return task_info
            
        except Exception as e:
            error_msg = str(e)
            
            # Handle specific error types
            if 'NotRegistered' in error_msg:
                return {
                    'state': 'FAILURE',
                    'result': {
                        'error': 'Task not found or not registered with worker. Make sure the Celery worker is running.',
                        'success': False
                    }
                }
            elif 'Redis' in error_msg or 'Connection refused' in error_msg:
                return {
                    'state': 'FAILURE', 
                    'result': {
                        'error': 'Redis connection failed. Check if Redis server is running.',
                        'success': False
                    }
                }
            else:
                return {
                    'state': 'FAILURE',
                    'result': {
                        'error': f"Failed to get task status: {error_msg}",
                        'success': False
                    }
                }
    @rate_limit("batch_processing", limit=5, window=600)  # 5 batch operations per 10 minutes
    def process_all_resumes(self, files: List[Any]) -> List[Dict[str, Any]]:
        """
        Process all resumes for bulk preview/generation (no email).
        Returns a list of preview/result dicts for each file.
        """
        results = []
        for file in files:
            # Try to get file name and text from session state if available
            file_name = getattr(file, 'name', 'Resume')
            resume_inputs = st.session_state.get('resume_inputs', {})
            file_data = resume_inputs.get(file_name, {})
            text = file_data.get('text', '')
            manual_text = file_data.get('manual_text', '')
            # Use PreviewGenerator for consistent preview output
            try:
                from core.resume_processor import PreviewGenerator
                previewer = PreviewGenerator()
                preview_result = previewer.generate_preview(file, text, manual_text)
                preview_result['file_name'] = file_name
                results.append(preview_result)
            except Exception as e:
                results.append({'file_name': file_name, 'success': False, 'error': str(e)})
        return results
    """Main processor for individual resume operations."""
    
    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
        self.file_processor = FileProcessor()
        # Use enhanced robust processor for error recovery
        self.robust_processor = RobustResumeProcessor()
        # Use distributed cache instead of local cache
        self.cache_manager = get_distributed_cache_manager()
        # Progress tracking
        self.progress_tracker = get_progress_tracker()
        # Legacy support - Advanced cache: (filename, text, manual_text) -> result
        self._result_cache: Dict[str, Dict[str, Any]] = {}
        self._cache_lock = threading.Lock()
        self._audit_log = []
        self._task_queue = queue.Queue()
        self._worker_thread = threading.Thread(target=self._worker_loop, daemon=True)
        self._worker_thread.start()

    def _worker_loop(self):
        while True:
            try:
                task, args, kwargs = self._task_queue.get()
                try:
                    task(*args, **kwargs)
                except Exception as e:
                    logger.error(f"Background task error: {e}")
            except Exception:
                pass
        
    @record_metrics("resume_processing")
    @file_processing_circuit_breaker
    @cached_processing(ttl=3600)
    def process_single_resume(
        self, 
        file_data: Dict[str, Any], 
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Process a single resume with enhanced error recovery, caching, and monitoring.
        
        Args:
            file_data: Dictionary containing file information and settings
            progress_callback: Optional callback for progress updates
            
        Returns:
            Result dictionary with processing status and data
        """
        start_time = time.time()
        cache_key = f"{file_data['filename']}|{file_data.get('text','')}|{file_data.get('manual_text','')}"
        with self._cache_lock:
            if cache_key in self._result_cache:
                logger.info(f"[CACHE] Hit for {file_data['filename']}")
                return self._result_cache[cache_key]
        try:
            filename = file_data['filename']
            file_obj = file_data['file']
            raw_text = file_data['text']
            user = None
            try:
                user = st.session_state.get('user_id', 'unknown')
            except Exception:
                user = 'unknown'
            audit_logger.log(
                action="process_single_resume",
                user=user,
                details={"filename": filename},
                status="started"
            )
            if progress_callback:
                progress_callback(f"Parsing tech stacks for {filename}...")
            # Parse tech stacks and points (now backed by LRU cache in parser)
            manual_text = file_data.get('manual_text', '')
            
            # If there's manual text, use it; otherwise use parsed tech stacks
            if manual_text.strip():
                # For manual text, create a simple structure
                selected_points = [manual_text.strip().split('\n')]
                tech_stacks_used = ['Manual Entry']
            else:
                selected_points, tech_stacks_used = self.text_parser.parse_tech_stacks(raw_text)
            
            if not selected_points or not tech_stacks_used:
                result = {
                    'success': False,
                    'error': ERROR_MESSAGES["no_tech_stacks"].format(filename=filename),
                    'filename': filename
                }
                audit_logger.log(
                    action="process_single_resume",
                    user=user,
                    details={"filename": filename},
                    status="failed",
                    error=result['error']
                )
                with self._cache_lock:
                    self._result_cache[cache_key] = result
                return result
            if progress_callback:
                progress_callback(f"Processing document for {filename}...")
            # Load and process document
            doc = Document(file_obj)
            projects_data = self.doc_processor.project_detector.find_projects_and_responsibilities(doc)
            if not projects_data:
                result = {
                    'success': False,
                    'error': ERROR_MESSAGES["no_projects"].format(filename=filename),
                    'filename': filename
                }
                audit_logger.log(
                    action="process_single_resume",
                    user=user,
                    details={"filename": filename},
                    status="failed",
                    error=result['error']
                )
                with self._cache_lock:
                    self._result_cache[cache_key] = result
                return result
            # Convert to structured format
            projects = []
            for i, project_tuple in enumerate(projects_data):
                if len(project_tuple) != 3:
                    logger.error(f"Malformed project tuple at index {i}: {project_tuple}")
                    continue
                title, start_idx, end_idx = project_tuple
                projects.append({
                    'title': title,
                    'index': i,
                    'responsibilities_start': start_idx,
                    'responsibilities_end': end_idx
                })
            # Distribute and add points using round-robin logic
            distribution_result = self.doc_processor.point_distributor.distribute_points_to_projects(
                projects, (selected_points, tech_stacks_used)
            )
            if not distribution_result['success']:
                with self._cache_lock:
                    self._result_cache[cache_key] = distribution_result
                return distribution_result
            # Add points to each project with mixed tech stacks
            total_added = 0
            # Sort projects by insertion point to process them in order
            sorted_projects = sorted(distribution_result['distribution'].items(),
                                   key=lambda x: x[1]['insertion_point'])
            # Keep track of how many paragraphs we've added to adjust insertion points
            paragraph_offset = 0
            for project_title, project_info in sorted_projects:
                # Adjust insertion point based on previous additions
                adjusted_project_info = project_info.copy()
                adjusted_project_info['insertion_point'] += paragraph_offset
                if 'responsibilities_end' in adjusted_project_info:
                    adjusted_project_info['responsibilities_end'] += paragraph_offset
                added = self.doc_processor.add_points_to_project(doc, adjusted_project_info)
                total_added += added
                # Update the offset for subsequent projects
                paragraph_offset += added
            points_added = total_added
            if progress_callback:
                progress_callback(f"Saving document for {filename}...")
            # Save to buffer
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            result = {
                'success': True,
                'filename': filename,
                'buffer': output_buffer.getvalue(),
                'tech_stacks': tech_stacks_used,
                'points_added': points_added,
                'email_data': self._extract_email_data(file_data)
            }
            audit_logger.log(
                action="process_single_resume",
                user=user,
                details={"filename": filename},
                status="success"
            )
            with self._cache_lock:
                self._result_cache[cache_key] = result
            return result
        except Exception as e:
            # Retry logic: up to 2 more times for transient errors
            for attempt in range(2):
                try:
                    time.sleep(0.5 * (attempt + 1))
                    return self.process_single_resume(file_data, progress_callback)
                except Exception:
                    continue
            logger.error(f"Error processing resume {file_data.get('filename','?')}: {e}")
            result = {'success': False, 'filename': file_data.get('filename', '?'), 'error': str(e)}
            audit_logger.log(
                action="process_single_resume",
                user=user if 'user' in locals() else 'unknown',
                details={"filename": file_data.get('filename', '?')},
                status="failed",
                error=str(e)
            )
            with self._cache_lock:
                self._result_cache[cache_key] = result
            return result
    
    def _extract_email_data(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract email configuration from file data.
        
        Args:
            file_data: File data dictionary
            
        Returns:
            Email configuration dictionary
        """
        return {
            'recipient': file_data.get('recipient_email', ''),
            'sender': file_data.get('sender_email', ''),
            'password': file_data.get('sender_password', ''),
            'smtp_server': file_data.get('smtp_server', ''),
            'smtp_port': file_data.get('smtp_port', 465),
            'subject': file_data.get('email_subject', ''),
            'body': file_data.get('email_body', '')
        }


class BulkResumeProcessor:
    """Handles bulk processing of multiple resumes with parallel execution."""
    
    def __init__(self):
        self.resume_processor = ResumeProcessor()
        self.file_processor = FileProcessor()
        self.email_manager = get_email_manager()
        # Determine a sensible default for max workers based on CPU
        try:
            import os
            cpu_count = os.cpu_count() or 4
            # Use min to avoid oversubscription; ThreadPool benefits from I/O overlap but docx is CPU-heavy
            self._default_workers = max(2, min(8, cpu_count))
        except Exception:
            self._default_workers = 4
    
    def process_resumes_bulk(
        self,
        files_data: List[Dict[str, Any]],
        max_workers: int = None,
        progress_callback: Optional[Callable] = None,
        use_process_pool: Optional[bool] = None,
        memory_limit_mb: int = 1024,
        ui_update_interval: float = 0.5
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Process multiple resumes in parallel.
        
        Args:
            files_data: List of file data dictionaries
            max_workers: Maximum number of parallel workers
            progress_callback: Optional progress callback
            
        Returns:
            Tuple of (processed_resumes, failed_resumes)
        """
        import time, psutil, os
        processed_resumes = []
        failed_resumes = []
        last_ui_update = time.time()
        if max_workers is None:
            cpu_count = os.cpu_count() or 4
            # For CPU-bound, use cpu_count; for I/O-bound, allow more
            max_workers = min(max(cpu_count, 4), len(files_data), 16)
        if max_workers < 1:
            max_workers = 1

        if use_process_pool is None:
            use_process_pool = len(files_data) >= 3 and (os.cpu_count() or 2) > 2

        # Ensure all files have proper names before processing
        optimized_files_data = []
        for file_data in files_data:
            file_data['file'] = self.file_processor.ensure_file_has_name(
                file_data['file'], file_data['filename']
            )
            optimized_files_data.append(file_data)

        def memory_check():
            try:
                process = psutil.Process(os.getpid())
                mem_mb = process.memory_info().rss / 1024 / 1024
                return mem_mb > memory_limit_mb
            except Exception:
                return False

        def throttled_progress_callback(msg):
            nonlocal last_ui_update
            now = time.time()
            if progress_callback and (now - last_ui_update > ui_update_interval):
                progress_callback(msg)
                last_ui_update = now

        if use_process_pool:
            # Only send minimal data: file path, not bytes
            # If file objects are not file paths, fallback to bytes
            payloads = []
            for file_data in optimized_files_data:
                file_obj = file_data['file']
                if hasattr(file_obj, 'name') and os.path.exists(file_obj.name):
                    payloads.append({
                        'filename': file_data['filename'],
                        'file_path': file_obj.name,
                        'text': file_data['text'],
                        'recipient_email': file_data.get('recipient_email', ''),
                        'sender_email': file_data.get('sender_email', ''),
                        'sender_password': file_data.get('sender_password', ''),
                        'smtp_server': file_data.get('smtp_server', ''),
                        'smtp_port': file_data.get('smtp_port', 465),
                        'email_subject': file_data.get('email_subject', ''),
                        'email_body': file_data.get('email_body', ''),
                    })
                else:
                    # fallback to bytes
                    file_bytes = file_obj.getvalue() if hasattr(file_obj, 'getvalue') else file_obj.read()
                    payloads.append({
                        'filename': file_data['filename'],
                        'file_content': file_bytes,
                        'text': file_data['text'],
                        'recipient_email': file_data.get('recipient_email', ''),
                        'sender_email': file_data.get('sender_email', ''),
                        'sender_password': file_data.get('sender_password', ''),
                        'smtp_server': file_data.get('smtp_server', ''),
                        'smtp_port': file_data.get('smtp_port', 465),
                        'email_subject': file_data.get('email_subject', ''),
                        'email_body': file_data.get('email_body', ''),
                    })

            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {executor.submit(_process_single_resume_worker, p): p['filename'] for p in payloads}
                for future in concurrent.futures.as_completed(future_to_file):
                    filename = future_to_file[future]
                    try:
                        result = future.result()
                        if result['success']:
                            processed_resumes.append(result)
                        else:
                            failed_resumes.append(result)
                        throttled_progress_callback(f"Processed {filename}")
                        if memory_check():
                            break
                    except Exception as e:
                        failed_resumes.append({'success': False, 'filename': filename, 'error': str(e)})
        else:
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                future_to_file = {
                    executor.submit(self.resume_processor.process_single_resume, file_data, throttled_progress_callback):
                    file_data['filename'] for file_data in optimized_files_data
                }
                for future in concurrent.futures.as_completed(future_to_file):
                    filename = future_to_file[future]
                    try:
                        result = future.result()
                        if result['success']:
                            processed_resumes.append(result)
                        else:
                            failed_resumes.append(result)
                        throttled_progress_callback(f"Processed {filename}")
                        if memory_check():
                            break
                    except Exception as e:
                        failed_resumes.append({'success': False, 'filename': filename, 'error': str(e)})

        self.file_processor.cleanup_memory()
        return processed_resumes, failed_resumes
    
    def send_batch_emails(
        self, 
        processed_resumes: List[Dict[str, Any]], 
        progress_callback: Optional[Callable] = None
    ) -> List[Dict[str, Any]]:
        """
        Send emails for processed resumes.
        
        Args:
            processed_resumes: List of processed resume data
            progress_callback: Optional progress callback
            
        Returns:
            List of email sending results
        """
        return self.email_manager.send_batch_emails(processed_resumes, progress_callback)


class PreviewGenerator:
    """Generates preview of resume changes without modifying the original."""
    
    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
    
    def generate_preview(
        self, 
        file_obj, 
        input_text: str, 
        manual_text: str = ""
    ) -> Dict[str, Any]:
        """
        Generate a preview of changes that will be made to the resume.
        
        Args:
            file_obj: Resume file object
            input_text: Tech stack input text
            manual_text: Optional manual points text
            
        Returns:
            Preview result dictionary
        """
        try:
            # Parse input text
            selected_points, tech_stacks_used = parse_input_text(input_text, manual_text)
            
            if not selected_points or not tech_stacks_used:
                return {
                    'success': False,
                    'error': 'Could not parse tech stacks from input'
                }
            
            # Load document and find projects
            doc = Document(file_obj)
            projects_data = self.doc_processor.project_detector.find_projects_and_responsibilities(doc)
            
            if not projects_data:
                return {
                    'success': False,
                    'error': 'No projects with Responsibilities sections found'
                }
            
            # Convert to structured format
            projects = []
            for i, (title, start_idx, end_idx) in enumerate(projects_data):
                projects.append({
                    'title': title,
                    'index': i,
                    'responsibilities_start': start_idx,
                    'responsibilities_end': end_idx
                })
            
            # Create preview document copy
            temp_buffer = BytesIO()
            doc.save(temp_buffer)
            temp_buffer.seek(0)
            
            preview_doc = Document(temp_buffer)
            preview_projects_data = self.doc_processor.project_detector.find_projects_and_responsibilities(preview_doc)
            
            # Convert to structured format with defensive check
            preview_projects = []
            for i, project_tuple in enumerate(preview_projects_data):
                if len(project_tuple) == 3:
                    title, start_idx, end_idx = project_tuple
                    preview_projects.append({
                        'title': title,
                        'index': i,
                        'responsibilities_start': start_idx,
                        'responsibilities_end': end_idx
                    })
                else:
                    logger.error(f"Malformed project tuple at index {i}: {project_tuple}")
                    continue
            
            # Apply changes to preview using round-robin logic
            distribution_result = self.doc_processor.point_distributor.distribute_points_to_projects(
                preview_projects, (selected_points, tech_stacks_used)
            )
            if not distribution_result['success']:
                return {'success': False, 'error': distribution_result['error']}
            
            # Add points to each project with mixed tech stacks for preview
            total_added = 0
            project_points_mapping = {}
            
            # Sort projects by insertion point to process them in order
            sorted_projects = sorted(distribution_result['distribution'].items(),
                                   key=lambda x: x[1]['insertion_point'])
            
            # Keep track of how many paragraphs we've added to adjust insertion points
            paragraph_offset = 0
            
            for project_title, project_info in sorted_projects:
                # Store which points are added to which project
                project_points = {}
                for tech_name, points in project_info['mixed_tech_stacks'].items():
                    project_points[tech_name] = points
                
                project_points_mapping[project_title] = project_points
                
                # Adjust insertion point based on previous additions
                adjusted_project_info = project_info.copy()
                adjusted_project_info['insertion_point'] += paragraph_offset
                if 'responsibilities_end' in adjusted_project_info:
                    adjusted_project_info['responsibilities_end'] += paragraph_offset
                
                # Add points to the document
                added = self.doc_processor.add_points_to_project(preview_doc, adjusted_project_info)
                total_added += added
                
                # Update the offset for subsequent projects
                paragraph_offset += added
            
            points_added = total_added
            
            # Generate preview content
            preview_content = self._extract_preview_content(preview_doc)
            
            return {
                'success': True,
                'points_added': points_added,
                'tech_stacks_used': tech_stacks_used,
                'selected_points': selected_points,
                'preview_content': preview_content,
                'preview_doc': preview_doc,
                'projects_count': len(projects),
                'project_points_mapping': project_points_mapping
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_preview_content(self, doc: Document) -> str:
        """
        Extract readable content from document for preview.
        
        Args:
            doc: Document to extract content from
            
        Returns:
            Text content of the document
        """
        content_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)
        
        return "\n\n".join(content_parts)


class ResumeManager:
    def process_all_resumes(self, files: List[Any]) -> List[Dict[str, Any]]:
        """Process all resumes for bulk preview/generation (no email)."""
        return self.resume_processor.process_all_resumes(files)
    """Main manager class that coordinates all resume processing operations."""
    
    def __init__(self):
        self.resume_processor = ResumeProcessor()
        self.bulk_processor = BulkResumeProcessor()
        self.preview_generator = PreviewGenerator()
        self.email_manager = get_email_manager()
    
    def process_single_resume(
        self, 
        file_data: Dict[str, Any], 
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """Process a single resume."""
        return self.resume_processor.process_single_resume(file_data, progress_callback)
    
    def process_bulk_resumes(
        self, 
        files_data: List[Dict[str, Any]], 
        max_workers: int = None,
        progress_callback: Optional[Callable] = None
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process multiple resumes in parallel."""
        return self.bulk_processor.process_resumes_bulk(files_data, max_workers, progress_callback)
    
    # Async integration (Celery) delegations
    def process_single_resume_async(self, file_data: Dict[str, Any]):
        """Submit a resume processing job to Celery and return the AsyncResult."""
        return self.resume_processor.process_single_resume_async(file_data)

    def get_async_result(self, task_id: str) -> Dict[str, Any]:
        """Fetch status/result for an async Celery task by ID."""
        return self.resume_processor.get_async_result(task_id)

    def generate_preview(
        self, 
        file_obj, 
        input_text: str, 
        manual_text: str = ""
    ) -> Dict[str, Any]:
        """Generate preview of resume changes."""
        return self.preview_generator.generate_preview(file_obj, input_text, manual_text)
    
    def send_batch_emails(
        self, 
        processed_resumes: List[Dict[str, Any]], 
        progress_callback: Optional[Callable] = None
    ) -> List[Dict[str, Any]]:
        """Send batch emails for processed resumes."""
        return self.bulk_processor.send_batch_emails(processed_resumes, progress_callback)
    
    def send_single_email(
        self, 
        smtp_server: str,
        smtp_port: int,
        sender_email: str,
        sender_password: str,
        recipient_email: str,
        subject: str,
        body: str,
        attachment_data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """Send a single email."""
        return self.email_manager.send_single_email(
            smtp_server, smtp_port, sender_email, sender_password,
            recipient_email, subject, body, attachment_data, filename
        )
    
    def validate_email_config(self, email_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate email configuration."""
        return self.email_manager.validate_email_config(email_data)
    
    def cleanup(self) -> None:
        """Clean up resources and connections."""
        self.email_manager.close_all_connections()


# Global resume manager instance
resume_manager = ResumeManager()


@st.cache_resource
def get_resume_manager(_version: str = "v2.2") -> ResumeManager:
    """
    Get the global resume manager instance.
    
    Returns:
        ResumeManager instance
    """
    # Force cache refresh by creating new instance
    return ResumeManager()


